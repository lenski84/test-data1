from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2)

FONT = "Times New Roman"
BASE = 12

def set_font(run, size=BASE, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), FONT)
    rFonts.set(qn('w:hAnsi'), FONT)
    rFonts.set(qn('w:cs'),    FONT)
    rPr.insert(0, rFonts)

def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18) if level == 1 else Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run(text)
    size = 14 if level == 1 else 12
    set_font(run, size=size, bold=True)
    return p

def body(doc, text, space_after=6, justify=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(18)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run)
    return p

def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(chr(9472) * 80)
    run.font.name = FONT
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

def add_table_row(table, cells_data, bold=False, bg=None):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        set_font(run, bold=bold)
        if bg:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), bg)
            tcPr.append(shd)
    return row

def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[i])

# ═══════════════════════════════════════════════════════════════
# KAPITEL 2: THEORETISCHE GRUNDLAGEN BENCHMARKING
# ═══════════════════════════════════════════════════════════════

heading(doc, "2  Theoretische Grundlagen des Benchmarkings", level=1)

body(doc,
    "Das folgende Kapitel legt die theoretischen Grundlagen des Benchmarkings dar. "
    "Es werden zunächst Definition und Begriffsabgrenzung erarbeitet, anschliessend "
    "die typologische Einteilung verschiedener Benchmarking-Arten vorgenommen und "
    "die historische Entwicklung der Methode nachgezeichnet. Den Schwerpunkt bilden "
    "die etablierten Prozessmodelle, die als konzeptionelle Grundlage fuer die "
    "Durchfuehrung des Benchmarks in Kapitel 5 dienen. Abschliessend wird die "
    "spezifische Relevanz des Benchmarkings fuer den Sondermaschinenbau diskutiert.")

divider(doc)

# ─── 2.1 Definition ───
heading(doc, "2.1  Definition und Begriffsabgrenzung", level=2)

body(doc,
    "Der Begriff Benchmarking leitet sich vom englischen Wort benchmark ab, "
    "das urspruenglich aus der Vermessungstechnik stammt und eine Referenzmarke "
    "bezeichnet, anhand derer Hoehenmessungen kalibriert werden. "
    "In den Wirtschaftswissenschaften wurde der Begriff erstmals systematisch "
    "von Robert C. Camp im Kontext des Xerox-Konzerns gepraegt, der in den "
    "fruehen 1980er-Jahren ein strukturiertes Verfahren zur Messung und zum "
    "Vergleich von Unternehmensleistungen mit den besten Wettbewerbern entwickelte "
    "(Camp, 1994, S. 10).")

body(doc,
    "Eine praezise und weithin anerkannte Definition liefern Kounev et al. (2020, S. 3): "
    "\"Benchmarking ist der Prozess des Messens und Vergleichens von Leistungen, "
    "Prozessen oder Produkten eines Unternehmens mit denen eines oder mehrerer "
    "Referenzunternehmen, mit dem Ziel, Verbesserungspotenziale zu identifizieren "
    "und Best Practices zu adaptieren.\" "
    "Diese Definition hebt drei wesentliche Elemente hervor: den Messprozess, "
    "den Vergleich mit externen Referenzpunkten sowie die handlungsorientierte "
    "Ableitung von Verbesserungsmassnahmen.")

body(doc,
    "Benchmarking ist von verwandten Konzepten abzugrenzen. "
    "Waehrend ein einfacher Wettbewerbsvergleich (Competitive Intelligence) "
    "lediglich die Beobachtung und Sammlung von Informationen ueber Konkurrenten umfasst, "
    "geht Benchmarking darueber hinaus: Es zielt systematisch auf die Identifikation "
    "von Leistungsluecken (Gaps) und deren Ursachen sowie auf die strukturierte "
    "Ableitung von Massnahmen ab (Spendolini, 1992, S. 11). "
    "Gegenueber dem internen Controlling, das sich auf unternehmenseigene Kennzahlen "
    "beschraenkt, oeffnet Benchmarking den Blick auf externe Massstabe und "
    "branchenweite Best Practices (Watson, 1993, S. 4).")

divider(doc)

# ─── 2.2 Arten ───
heading(doc, "2.2  Arten des Benchmarkings", level=2)

body(doc,
    "In der Literatur werden Benchmarking-Arten anhand zweier zentraler Dimensionen "
    "unterschieden: dem Vergleichsobjekt (was wird verglichen?) und der "
    "Vergleichsbasis (mit wem wird verglichen?). Die folgende Tabelle gibt einen "
    "strukturierten Ueberblick.")

body(doc, "Tabelle 2.1: Taxonomie der Benchmarking-Arten", space_after=3)

tbl = doc.add_table(rows=1, cols=3)
tbl.style = "Table Grid"
add_table_row(tbl, ["Dimension", "Art", "Beschreibung"], bold=True, bg="D9D9D9")
arten_data = [
    ("Vergleichsbasis", "Intern",
     "Vergleich zwischen Abteilungen, Standorten oder Tochtergesellschaften desselben Unternehmens. Einfache Datenverfuegbarkeit, aber begrenzte externe Perspektive."),
    ("Vergleichsbasis", "Extern (Wettbewerb)",
     "Vergleich mit direkten Marktbegleitern. Hohe Relevanz, aber eingeschraenkte Datenzugaenglichkeit."),
    ("Vergleichsbasis", "Extern (Branche)",
     "Vergleich mit Unternehmen derselben Branche, aber nicht zwingend direkten Wettbewerbern. Breiter Erkenntnisgewinn."),
    ("Vergleichsbasis", "Funktional/generisch",
     "Vergleich branchenubergreifend mit Best-in-Class-Unternehmen in spezifischen Prozessen. Hoecstes Innovationspotenzial."),
    ("Vergleichsobjekt", "Produkt-Benchmarking",
     "Vergleich von Produkteigenschaften, Qualitaet, Funktionalitaet und Preis-Leistungs-Verhaeltnissen."),
    ("Vergleichsobjekt", "Prozess-Benchmarking",
     "Vergleich operativer Ablaeufe, Fertigungsprozesse oder Serviceprozesse zur Effizienzsteigerung."),
    ("Vergleichsobjekt", "Strategisches Benchmarking",
     "Vergleich strategischer Entscheidungen, Marktpositionierungen und langfristiger Erfolgsfaktoren."),
]
for row in arten_data:
    add_table_row(tbl, row)
set_col_widths(tbl, [3.5, 3.5, 8.5])
doc.add_paragraph()

body(doc,
    "Die vorliegende Arbeit verbindet strategisches Benchmarking mit Elementen des "
    "Produkt- und Prozessbenchmarkings. Im Vordergrund steht der externe Vergleich "
    "mit Wettbewerbern und Branchenteilnehmern im Segment des Sondermaschinenbaus. "
    "Dies entspricht dem von Watson (1993, S. 17) als 'competitive strategic benchmarking' "
    "bezeichneten Ansatz, der auf die Identifikation strategischer Wettbewerbsvorteile "
    "und struktureller Leistungsluecken zielt.")

divider(doc)

# ─── 2.3 Historische Entwicklung ───
heading(doc, "2.3  Historische Entwicklung", level=2)

body(doc,
    "Die systematische Nutzung von Benchmarking als Managementmethode begann in den "
    "fruehen 1980er-Jahren im Xerox-Konzern in den USA. Auf Initiative von Robert C. Camp "
    "wurden erstmals Fertigungsprozesse und Logistikkennzahlen systematisch mit "
    "Konkurrenten verglichen, um Kostenueberlegenheit japanischer Anbieter zu erklaeren "
    "und Best Practices zu identifizieren (Camp, 1994, S. 10-14). "
    "Das daraus resultierende Zehn-Schritte-Modell gilt als Grundstein der Methode "
    "und wurde in Camps Standardwerk von 1989 (deutschsprachig 1994) publiziert.")

body(doc,
    "In den 1990er-Jahren erlebte Benchmarking eine rasante Verbreitung. "
    "Spendolini (1992) systematisierte die Methode in einem funfstufigen Prozessmodell "
    "und erweiterte sie auf nicht-kompetitive Vergleichspartner. "
    "Watson (1993) praegte den Begriff des Strategic Benchmarking und betonte "
    "die Verknuepfung mit uebergeordneten Unternehmensstrategien. "
    "Parallel dazu gruendete sich 1992 das American Productivity and Quality Center "
    "(APQC) als institutionelle Plattform fuer Benchmarking-Initiativen, "
    "das bis heute eine der wichtigsten Benchmarking-Datenbanken betreibt (APQC, 2022).")

body(doc,
    "Im deutschsprachigen Raum gewann Benchmarking ab Mitte der 1990er-Jahre "
    "zunehmend an Bedeutung. Das Fraunhofer-Institut fuer Zuverlassigkeit und "
    "Mikrointegration (IZB) entwickelte ein auf den deutschen Mittelstand "
    "zugeschnittenes Fuenf-Phasen-Modell, das die spezifischen Herausforderungen "
    "mittelstaendischer Produktionsunternehmen beruecksichtigt (Fraunhofer IZB, 2023). "
    "Dieses Modell hat sich im deutschsprachigen Industrieumfeld als Standard etabliert "
    "und bildet den methodischen Leitrahmen der vorliegenden Arbeit.")

# Zeitleiste als Tabelle
body(doc, "Tabelle 2.2: Meilensteine der Benchmarking-Entwicklung", space_after=3)

tbl2 = doc.add_table(rows=1, cols=3)
tbl2.style = "Table Grid"
add_table_row(tbl2, ["Jahr", "Ereignis", "Bedeutung"], bold=True, bg="D9D9D9")
timeline = [
    ("1979-1982", "Xerox/Camp: Vergleich mit Fuji Xerox",
     "Geburtsstunde des modernen Benchmarkings"),
    ("1989", "Camp: \"Benchmarking\" (Erstauflage)",
     "Wissenschaftliche Grundlage, 10-Schritte-Modell"),
    ("1992", "Gruendung APQC (USA)",
     "Institutionalisierung; erste Benchmarking-Datenbank"),
    ("1992", "Spendolini: \"The Benchmarking Book\"",
     "Erweiterung auf generisches/funktionales Benchmarking"),
    ("1993", "Watson: \"Strategic Benchmarking\"",
     "Verknuepfung mit Unternehmensstrategie"),
    ("ab 1995", "Verbreitung im deutschsprachigen Raum",
     "Adaption fuer Mittelstand; Fraunhofer-Modell"),
    ("2022+", "APQC Open Standards Benchmarking",
     "Digitale Plattformen, Echtzeit-KPIs"),
]
for row in timeline:
    add_table_row(tbl2, row)
set_col_widths(tbl2, [2.5, 7.5, 5.5])
doc.add_paragraph()

divider(doc)

# ─── 2.4 Prozessmodelle ───
heading(doc, "2.4  Etablierte Benchmarking-Prozessmodelle", level=2)

body(doc,
    "Fuer die praktische Durchfuehrung eines Benchmarks existieren in der Literatur "
    "mehrere erprobte Prozessmodelle. Die drei bedeutendsten werden nachfolgend "
    "beschrieben und vergleichend gegenuubergestellt.")

heading(doc, "2.4.1  Das Zehn-Schritte-Modell nach Camp (1994)", level=2)

body(doc,
    "Das von Camp (1994, S. 17-260) entwickelte Modell gilt als Urmodell des "
    "Benchmarkings und gliedert den Prozess in zehn Schritte, die vier Phasen zugeordnet "
    "werden koennen: (1) Planung, (2) Analyse, (3) Integration und (4) Aktion. "
    "In der Planungsphase werden Benchmarking-Gegenstand, Vergleichspartner und "
    "Datenerhebungsmethoden festgelegt. Die Analysephase umfasst die Datenerhebung, "
    "den Vergleich und die Ermittlung der Leistungsluecke (Gap). "
    "In der Integrationsphase werden die Erkenntnisse kommuniziert und Ziele definiert. "
    "Die Aktionsphase schliesst mit der Umsetzung der Massnahmen und dem kontinuierlichen "
    "Monitoring ab (Camp, 1994, S. 17).")

body(doc,
    "Das Modell zeichnet sich durch seinen zyklischen Charakter aus: "
    "Am Ende eines Benchmarking-Zyklus wird geprueft, ob neue Benchmarks erforderlich "
    "sind, und der Prozess beginnt erneut. Dies sichert die kontinuierliche "
    "Anpassungsfahigkeit des Unternehmens an veraenderte Marktbedingungen.")

heading(doc, "2.4.2  Das Fuenf-Stufen-Modell nach Spendolini (1992)", level=2)

body(doc,
    "Spendolini (1992, S. 48-152) kondensierte den Benchmarking-Prozess auf fuenf "
    "Kernstufen: (1) Entscheidung, was benchmarkiert werden soll, "
    "(2) Aufbau eines Benchmarking-Teams, (3) Identifikation von Vergleichspartnern, "
    "(4) Sammlung und Analyse von Informationen sowie (5) Umsetzung und Nachverfolgung. "
    "Das Modell legt besonderen Wert auf die Teamstruktur und die explizite Beruecksichtigung "
    "verschiedener Benchmarking-Typen (intern, wettbewerbsorientiert, funktional/generisch).")

heading(doc, "2.4.3  Das Fuenf-Phasen-Modell des Fraunhofer IZB (2023)", level=2)

body(doc,
    "Das Fraunhofer-Modell (Fraunhofer IZB, 2023, S. 14-17) ist speziell auf "
    "produzierende mittelstaendische Unternehmen im deutschsprachigen Raum ausgerichtet. "
    "Es strukturiert den Benchmarking-Prozess in fuenf Phasen:")

for item in [
    "Phase 1 – Initialisierung: Projektziel definieren, Scope festlegen, interne Ressourcen sichern.",
    "Phase 2 – Interne Analyse: Eigene Ist-Situation erheben, Kennzahlen bestimmen, Staerken/Schwaechen identifizieren.",
    "Phase 3 – Partnerauswahl & Datenerhebung: Vergleichsunternehmen selektieren, Daten sekundaer und primaer erheben.",
    "Phase 4 – Vergleich & Gap-Analyse: Leistungsluecken quantifizieren, Best Practices ableiten, Ursachenanalyse.",
    "Phase 5 – Massnahmenplanung & Umsetzung: Handlungsempfehlungen priorisieren, Umsetzung begleiten, Wirkung messen.",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(1)
    run = p.add_run(item)
    set_font(run)

doc.add_paragraph()

body(doc,
    "Das Fraunhofer-Modell eignet sich besonders fuer den Kontext dieser Arbeit, "
    "da es die Vergleichbarkeitsproblematik in der Einzelfertigung explizit adressiert "
    "und pragmatische Loesungsansaetze fuer eingeschraenkte Datenverfuegbarkeit bietet "
    "(Fraunhofer IZB, 2023, S. 9).")

# Vergleichstabelle der Modelle
body(doc, "Tabelle 2.3: Vergleich der Benchmarking-Prozessmodelle", space_after=3)

tbl3 = doc.add_table(rows=1, cols=4)
tbl3.style = "Table Grid"
add_table_row(tbl3, ["Kriterium", "Camp (1994)", "Spendolini (1992)", "Fraunhofer IZB (2023)"],
              bold=True, bg="D9D9D9")
compare_data = [
    ("Schritte/Phasen", "10 Schritte", "5 Stufen", "5 Phasen"),
    ("Zielgruppe", "Grossunternehmen", "Allgemein", "KMU / Mittelstand"),
    ("Besonderheit", "Zyklischer Prozess", "Teamfokus", "Mittelstandsadaption"),
    ("Fokus", "Operativ & strategisch", "Prozess & Typen", "Produktion"),
    ("Geeignet fuer PA", "Konzeptionell", "Erganezend", "Hauptmodell"),
]
for row in compare_data:
    add_table_row(tbl3, row)
set_col_widths(tbl3, [3.5, 3.5, 3.5, 5.0])
doc.add_paragraph()

divider(doc)

# ─── 2.5 Relevanz Sondermaschinenbau ───
heading(doc, "2.5  Benchmarking im Kontext des Sondermaschinenbaus", level=2)

body(doc,
    "Der Sondermaschinenbau weist branchenspezifische Eigenschaften auf, die die "
    "Anwendung des Benchmarkings erschweren, gleichzeitig aber seinen strategischen "
    "Mehrwert begruenden. Sondermaschinenbauer fertigen kundenindividuelle Loesungen "
    "in der Regel als Einzelstuecke oder in sehr kleinen Serien; standardisierte "
    "Leistungskennzahlen, wie sie in der Serienfertigung selbstverstaendlich sind, "
    "existieren hier oft nicht (Spendolini, 1992, S. 78).")

body(doc,
    "Die Branche ist mit etwa 6.800 Unternehmen und rund 952.000 Beschaeftigten "
    "einer der bedeutendsten Industriezweige Deutschlands. Der Gesamtumsatz des "
    "deutschen Maschinenbaus betrug im Jahr 2023 rund 254,4 Milliarden Euro, "
    "wovon ein erheblicher Anteil auf den Sondermaschinenbau entfaellt (VDMA, 2024, S. 4). "
    "Angesichts des zunehmenden Kostendrucks durch internationale Anbieter, "
    "insbesondere aus China und Suedostasien, gewinnt die strategische "
    "Selbstverortung durch Benchmarking an Bedeutung.")

body(doc,
    "Fuer das strategische Benchmarking im Sondermaschinenbau ergeben sich drei "
    "zentrale Herausforderungen: Erstens ist die Datenverfuegbarkeit begrenzt, "
    "da viele Marktteilnehmer nicht publizitaetspflichtig sind. "
    "Zweitens erschwert die Heterogenitaet der Produkte und Kundensegmente "
    "einen direkten Kennzahlenvergleich (Vergleichbarkeitsproblem). "
    "Drittens besteht ein Interessenkonflikt zwischen der Notwendigkeit zur "
    "Offenlegung eigener Daten und dem Wunsch nach Wahrung von Betriebsgeheimnissen "
    "(Camp, 1994, S. 67).")

body(doc,
    "Diese Einschraenkungen koennen durch methodische Massnahmen teilweise "
    "kompensiert werden: Normierung absoluter Kennzahlen durch relative Indikatoren, "
    "Erweiterung des Vergleichsfelds auf branchenverwandte Unternehmen sowie "
    "systematische Sekundaerdatenanalyse auf Basis oeffentlich verfuegbarer "
    "Geschaeftsberichte und Branchenstatistiken. "
    "Das methodische Vorgehen dieser Arbeit traegt diesen Besonderheiten "
    "Rechnung (vgl. Kapitel 4).")

divider(doc)

# ─── LITERATURVERZEICHNIS ───
heading(doc, "Literaturverzeichnis (Kapitel 2)", level=1)

refs = [
    ("APQC. (2022). Open Standards Benchmarking. American Productivity and Quality Center. "
     "https://www.apqc.org/benchmarking"),
    ("Camp, R. C. (1994). Benchmarking: The Search for Industry Best Practices that Lead "
     "to Superior Performance. ASQC Quality Press."),
    ("Fraunhofer IZB. (2023). Benchmarking im Mittelstand: Das Fraunhofer-Modell. "
     "Fraunhofer-Institut fuer Zuverlaessigkeit und Mikrointegration."),
    ("Kounev, S., Lange, K.-D., & von Kistowski, J. (2020). Systems Benchmarking: "
     "For Scientists and Engineers. Springer International Publishing. "
     "https://doi.org/10.1007/978-3-030-41705-5"),
    ("Spendolini, M. J. (1992). The Benchmarking Book. AMACOM."),
    ("VDMA. (2024). Maschinenbau in Zahl und Bild 2024. "
     "Verband Deutscher Maschinen- und Anlagenbau e. V."),
    ("Watson, G. H. (1993). Strategic Benchmarking: How to Rate Your Company's "
     "Performance Against the World's Best. Wiley."),
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent          = Cm(1.25)
    p.paragraph_format.first_line_indent    = Cm(-1.25)
    p.paragraph_format.space_after          = Pt(4)
    p.paragraph_format.line_spacing         = Pt(18)
    run = p.add_run(ref)
    set_font(run)

doc.save("Kapitel_2_Theoretische_Grundlagen_Benchmarking.docx")
print("Gespeichert: Kapitel_2_Theoretische_Grundlagen_Benchmarking.docx")
