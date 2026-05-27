from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# --- Seitenränder DIN A4 ---
sec = doc.sections[0]
sec.page_height = Cm(29.7)
sec.page_width  = Cm(21.0)
sec.left_margin   = Cm(3.0)
sec.right_margin  = Cm(2.5)
sec.top_margin    = Cm(2.5)
sec.bottom_margin = Cm(2.0)

# --- Style-Grundlage ---
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)
normal.paragraph_format.space_after  = Pt(6)
normal.paragraph_format.line_spacing = Pt(18)

# --- Hilfsfunktionen ---
def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.name  = 'Times New Roman'
        r.font.size  = Pt(14)
        r.font.bold  = True
        r.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(8)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.name  = 'Times New Roman'
        r.font.size  = Pt(13)
        r.font.bold  = True
        r.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    for r in p.runs:
        r.font.name   = 'Times New Roman'
        r.font.size   = Pt(12)
        r.font.bold   = True
        r.font.italic = False
        r.font.color.rgb = RGBColor(0, 0, 0)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    return p

def para(text, bold_start=None):
    """Blocksatz-Absatz; bold_start = String am Anfang fett setzen bis zum ersten ':'."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0)
    if bold_start and text.startswith(bold_start):
        rest = text[len(bold_start):]
        r = p.add_run(bold_start)
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r2 = p.add_run(rest)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(12)
    else:
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    return p

def table(headers, rows, caption=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    # Header
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.paragraphs[0].clear()
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Rows
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.paragraphs[0].clear()
            r = c.paragraphs[0].add_run(cell_text)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cp.paragraph_format.space_before = Pt(2)
        cp.paragraph_format.space_after  = Pt(8)
        for r in cp.runs:
            r.font.name   = 'Times New Roman'
            r.font.size   = Pt(10)
            r.font.italic = True
    else:
        doc.add_paragraph()
    return t

# ===========================================================
# KAPITEL 3
# ===========================================================

h1("3  Strategische Unternehmens- und Wettbewerbsanalyse")
para(
    "Das folgende Kapitel liefert die empirische und analytische Grundlage für den in "
    "Kapitel 5 durchgeführten Unternehmensvergleich. Zunächst wird Schaeffler SMB als "
    "Untersuchungsobjekt vorgestellt (3.1), bevor die relevante Branche systematisch "
    "analysiert wird (3.2). Daran schließt sich eine Wettbewerbsanalyse an (3.3), die "
    "mit einer internen Unternehmensbetrachtung (3.4) abgerundet wird."
)

# -----------------------------------------------------------
# 3.1
# -----------------------------------------------------------
h2("3.1  Unternehmensportrait Schaeffler SMB")

h3("3.1.1  Gründung, Rechtsform, Eigentümerstruktur und Standorte")
para(
    "Schaeffler SMB (Sondermaschinenbau) wurde im Jahr 1970 gegründet und ist als "
    "spezialisierte Einheit innerhalb der Schaeffler-Gruppe für die Entwicklung und "
    "Fertigung kundenspezifischer Sondermaschinen tätig. Die Schaeffler-Gruppe zählt zu "
    "den weltweit führenden Unternehmen in der Herstellung von Wälzlagern, "
    "Linearsystemen und Präzisionskomponenten für die Automobil- und Industrietechnik. "
    "Schaeffler SMB agiert dabei als kompetenter interner und externer Maschinenlieferant "
    "im Rahmen dieses globalen Industriekonzerns."
)
para("Rechtsform: [Platzhalter – bitte ergänzen, z. B. GmbH oder eigenständige Abteilung]")
para(
    "Eigentümerstruktur: Schaeffler SMB ist Teil der Schaeffler-Gruppe, die mehrheitlich "
    "im Familienbesitz der Gesellschafterfamilien Schaeffler-Thumann und Schaeffler gehalten "
    "wird. [Platzhalter – ggf. spezifizieren, ob SMB als eigenständige juristische Person "
    "oder als Geschäftsbereich geführt wird]"
)
para("Standorte: Schaeffler SMB ist weltweit tätig. [Platzhalter – Hauptstandort(e) und internationale Niederlassungen ergänzen]")

h3("3.1.2  Geschäftsmodell & Leistungsportfolio")
para(
    "Das Geschäftsmodell von Schaeffler SMB basiert auf der Entwicklung, Konstruktion und "
    "Fertigung von Sondermaschinen, die auf die spezifischen Produktions- und "
    "Prozessanforderungen der Kunden zugeschnitten sind. Im Gegensatz zu Herstellern von "
    "Standardmaschinen operiert Schaeffler SMB im projektgetriebenen Einzelauftragsgeschäft "
    "(Engineer-to-Order): Jede Maschine wird in enger Abstimmung mit dem Kunden konzipiert "
    "und realisiert. Das Leistungsportfolio umfasst neben der Maschinenkonstruktion "
    "typischerweise auch Montage, Inbetriebnahme, Schulung und Aftersales-Service."
)
para("[Platzhalter – spezifische Produktlinien, Technologiebereiche und Branchenschwerpunkte ergänzen]")

h3("3.1.3  Aktuelle Marktposition & Kennzahlen")
table(
    ["Kennzahl", "Wert"],
    [
        ["Gründungsjahr", "1970"],
        ["Mitarbeiter (weltweit)", "ca. 1.600"],
        ["Umsatz", "Hoher zweistelliger Millionenbereich (€)"],
        ["Rechtsform", "[Platzhalter]"],
        ["Hauptstandort", "[Platzhalter]"],
        ["Märkte", "International / weltweit"],
        ["Muttergesellschaft", "Schaeffler-Gruppe"],
    ],
    caption="Tabelle 3.1: Eckdaten Schaeffler SMB. Quelle: Eigene Darstellung.",
)
para(
    "Mit rund 1.600 Mitarbeiterinnen und Mitarbeitern und einem Umsatz im hohen "
    "zweistelligen Millionenbereich positioniert sich Schaeffler SMB als mittelgroßer, "
    "hochspezialisierter Anbieter im Sondermaschinenbau. Die Einbindung in die "
    "Schaeffler-Gruppe verleiht dem Unternehmen Zugang zu globalem Know-how, "
    "internationalen Produktionsnetzwerken und einem etablierten Kundenstamm in der "
    "Automobil- und Industrietechnik."
)

# -----------------------------------------------------------
# 3.2
# -----------------------------------------------------------
h2("3.2  Branchenanalyse")
para(
    "Die strategische Positionierung eines Unternehmens lässt sich nur im Kontext der "
    "relevanten Branche beurteilen. Nachfolgend wird der Markt für Sondermaschinenbau "
    "hinsichtlich Struktur, Größe, Trends und Wettbewerbsdynamik analysiert."
)

h3("3.2.1  Marktstruktur & Marktgröße")
para(
    "Der Sondermaschinenbau ist ein Teilsegment des deutschen Maschinenbaus, der insgesamt "
    "zur bedeutendsten Industriebranche Deutschlands zählt. Nach Angaben des Verbands "
    "Deutscher Maschinen- und Anlagenbau (VDMA) erzielte die Branche im Jahr 2024 einen "
    "Gesamtumsatz von 254,4 Milliarden Euro und beschäftigte rund 952.000 Personen in "
    "Betrieben mit mindestens 50 Mitarbeitenden (VDMA, 2025). Trotz eines Rückgangs "
    "gegenüber dem Rekordjahr 2023 (262,9 Mrd. €) bleibt der Maschinenbau das Rückgrat "
    "der deutschen Industrie."
)
para(
    "Der Sondermaschinenbau als Spezialsegment zeichnet sich durch eine ausgeprägte "
    "Fragmentierung aus: Einer kleinen Anzahl großer Systemlieferanten steht eine Vielzahl "
    "mittelständischer Nischenanbieter gegenüber, die sich auf spezifische Branchen oder "
    "Technologiebereiche spezialisiert haben. Die Exportquote ist mit 199,6 Milliarden Euro "
    "im Jahr 2024 außerordentlich hoch, was die internationale Wettbewerbsfähigkeit des "
    "Segments unterstreicht (VDMA, 2025)."
)
para(
    "[Platzhalter – sofern verfügbar: spezifische Marktvolumendaten für den "
    "Sondermaschinenbau ergänzen, z. B. VDMA-Fachabteilung oder einschlägige Marktstudien]"
)

h3("3.2.2  Markttrends & Treiber")
para(
    "Die Branche unterliegt einem tiefgreifenden strukturellen Wandel, der durch mehrere "
    "übergeordnete Trends getrieben wird (Grant, 2024, [S. X]; Whittington et al., 2024, [S. X]):"
)
para("Digitalisierung und Industrie 4.0: Die Integration von Sensorik, Konnektivität und "
     "Datenanalyse in Maschinen und Produktionssysteme verändert das Leistungsangebot im "
     "Sondermaschinenbau grundlegend. Kunden erwarten zunehmend Maschinen, die sich in "
     "übergeordnete MES- und ERP-Systeme integrieren lassen und Predictive-Maintenance-"
     "Funktionen bieten.", "Digitalisierung und Industrie 4.0:")
para("Automatisierung: Der anhaltende Trend zur Automatisierung manueller Produktionsprozesse "
     "steigert die Nachfrage nach Sondermaschinen, insbesondere in lohnkostenintensiven "
     "Branchen wie der Automobilindustrie.", "Automatisierung:")
para("Reshoring und Nearshoring: Geopolitische Unsicherheiten und Lieferkettenunterbrechungen "
     "veranlassen Unternehmen, Produktionskapazitäten zurück nach Europa zu verlagern. Dies "
     "schafft Nachfragepotenzial für neue Produktionsanlagen und damit für Sondermaschinen-"
     "hersteller.", "Reshoring und Nearshoring:")
para("Fachkräftemangel: Der demografisch bedingte Rückgang des Erwerbspersonenpotenzials "
     "trifft den Sondermaschinenbau besonders hart, da die Branche auf hochqualifizierte "
     "Ingenieure und Facharbeiter angewiesen ist (VDMA, 2025).", "Fachkräftemangel:")
para("Konsolidierungstendenzen: Wachsender Kostendruck und steigende Anforderungen an die "
     "digitale Kompetenz begünstigen Zusammenschlüsse und Übernahmen, insbesondere unter "
     "kleineren Anbietern (Grant, 2024, [S. X]).", "Konsolidierungstendenzen:")

h3("3.2.3  PESTEL-Analyse")
para(
    "Die PESTEL-Analyse ist ein etabliertes Instrument zur systematischen Erfassung der "
    "makroökonomischen Umweltfaktoren, die auf ein Unternehmen und seine Branche einwirken "
    "(Whittington et al., 2024, [S. X]). Sie gliedert sich in sechs Dimensionen: politische "
    "(Political), wirtschaftliche (Economic), sozio-kulturelle (Social), technologische "
    "(Technological), ökologische (Environmental) und rechtliche (Legal) Faktoren."
)
table(
    ["Faktor", "Ausprägung im Sondermaschinenbau", "Bewertung"],
    [
        ["Politisch (P)",
         "Handelspolitische Spannungen (v. a. USA–China) erschweren internationale "
         "Projektgeschäfte; Exportkontrollregularien für Dual-Use-Güter (AWG, EAR) "
         "erhöhen Compliance-Aufwand; Förderung von Industrieautomatisierung durch "
         "EU-Industriepolitik (IPCEI, Chips Act).",
         "Mittel\nRelevanz steigt durch geopolitische Fragmentierung"],
        ["Ökonomisch (E)",
         "Investitionszyklen der Abnehmerindustrien (v. a. Automobil) bestimmen die "
         "Nachfrage; höheres Zinsniveau (2022–2024) dämpft Investitionsbereitschaft; "
         "Energiepreise beeinflussen Produktionskosten und Wettbewerbsfähigkeit der Kunden.",
         "Hoch\nDirekte Wirkung auf Auftragseingang"],
        ["Sozio-kulturell (S)",
         "Akuter Fachkräftemangel bei MINT-Berufen; demografischer Wandel reduziert "
         "Arbeitskräfteangebot langfristig; steigende Ansprüche an Arbeitsbedingungen "
         "und Work-Life-Balance (VDMA, 2025).",
         "Hoch\nStrategisches Risiko für Personalverfügbarkeit"],
        ["Technologisch (T)",
         "Industrie 4.0, KI-Integration (ML für Qualitätsprüfung), digitale Zwillinge "
         "(Simulation vor Bau), Kollaborationsroboter (Cobots); Plattformisierung der "
         "Maschinensteuerung.",
         "Sehr hoch\nDisruptives Potenzial und Chancen"],
        ["Ökologisch (E)",
         "EU-Taxonomie und CSRD erhöhen Nachhaltigkeitsanforderungen; "
         "Energieeffizienzanforderungen an Maschinen steigen (Ökodesign-Verordnung); "
         "Kreislaufwirtschaftsanforderungen für Maschinenbaukomponenten.",
         "Mittel–Hoch\nWächst regulatorisch"],
        ["Legal (L)",
         "EU-Maschinenverordnung 2023/1230/EU (ersetzt Richtlinie 2006/42/EG); "
         "CE-Kennzeichnungspflicht; Produkthaftungsgesetz; DSGVO für vernetzte Maschinen; "
         "KI-Verordnung (AI Act) bei KI-gestützten Maschinen.",
         "Hoch\nCompliance-Aufwand steigt kontinuierlich"],
    ],
    caption=(
        "Tabelle 3.2: PESTEL-Analyse Sondermaschinenbau. "
        "Quelle: Eigene Darstellung in Anlehnung an Whittington et al. (2024, [S. X]); VDMA (2025)."
    ),
)

h3("3.2.4  Branchenstrukturanalyse nach Porter (Five Forces)")
para(
    "Das von Porter (1980) entwickelte Fünf-Kräfte-Modell analysiert die strukturellen "
    "Wettbewerbskräfte einer Branche und ermöglicht eine Einschätzung ihrer langfristigen "
    "Attraktivität (Porter, 1980; Porter, 2008). Die fünf Kräfte umfassen die Rivalität "
    "bestehender Wettbewerber, die Verhandlungsmacht von Lieferanten und Kunden sowie die "
    "Bedrohung durch Substitute und Neueintreter."
)
para(
    "Rivalität der bestehenden Wettbewerber (mittel–hoch): Der Sondermaschinenbau ist "
    "durch eine hohe Anbieterfragmentierung gekennzeichnet. Neben einigen wenigen "
    "Großanbietern konkurrieren zahlreiche mittelständische Spezialisten um Projekte. Da "
    "Aufträge in der Regel über Ausschreibungen vergeben werden, besteht ein erheblicher "
    "Preis- und Leistungswettbewerb. Technologisches Know-how schützt etablierte Anbieter "
    "jedoch vor schnellem Verdrängungswettbewerb (Porter, 2008, S. 79).",
    "Rivalität der bestehenden Wettbewerber (mittel–hoch):",
)
para(
    "Verhandlungsmacht der Lieferanten (mittel): Schlüsselkomponenten wie Antriebssysteme, "
    "Steuerungen und Sensorik stammen häufig von spezialisierten Zulieferern (z. B. Siemens, "
    "Bosch Rexroth, Fanuc). Diese besitzen durch ihre Marktstellung eine gewisse "
    "Lieferantenmacht. Gleichzeitig besteht für viele Standardkomponenten ein ausreichend "
    "breiter Lieferantenmarkt (Porter, 1980, [S. X]).",
    "Verhandlungsmacht der Lieferanten (mittel):",
)
para(
    "Verhandlungsmacht der Kunden (mittel–hoch): Im Sondermaschinenbau handelt es sich "
    "typischerweise um Großkunden aus der Automobil-, Elektronik- oder Prozessindustrie, "
    "die über erhebliche Nachfragemacht verfügen. Auf der anderen Seite entstehen durch die "
    "Spezifität der Maschinen hohe Wechselkosten, da der Lieferantenwechsel hohe "
    "Transaktionskosten verursacht (Porter, 1980, [S. X]).",
    "Verhandlungsmacht der Kunden (mittel–hoch):",
)
para(
    "Bedrohung durch Substitute (niedrig): Sondermaschinen sind per Definition auf "
    "spezifische Prozesse zugeschnitten und nicht direkt durch Standardmaschinen "
    "substituierbar. Eine gewisse Bedrohung besteht durch das Insourcing seitens großer "
    "Kunden sowie durch modulare Automatisierungsplattformen (Porter, 2008, S. 82).",
    "Bedrohung durch Substitute (niedrig):",
)
para(
    "Bedrohung durch Neueintreter (niedrig–mittel): Die Eintrittsbarrieren in den "
    "Sondermaschinenbau sind erheblich: Neben technologischem Know-how sind langjährige "
    "Kundenreferenzen, Prüfzertifizierungen und spezifisches Applikationswissen "
    "erforderlich. Dennoch können neue Technologieanbieter – etwa aus dem Bereich der "
    "Softwareautomatisierung oder Robotik – mit disruptiven Plattformansätzen in "
    "Teilsegmente eindringen (Grant, 2024, [S. X]).",
    "Bedrohung durch Neueintreter (niedrig–mittel):",
)
para(
    "Fazit zur Branchenattraktivität: In der Gesamtschau zeigt die Branchenstrukturanalyse "
    "eine moderat attraktive Branche. Hohe Eintrittsbarrieren und geringe Substituierbarkeit "
    "bieten Schutz für etablierte Anbieter. Gleichzeitig dämpfen die Verhandlungsmacht "
    "großer Kunden und der intensive Wettbewerb unter den Anbietern die Rentabilität. Die "
    "strukturelle Attraktivität hängt stark von der Fähigkeit des einzelnen Unternehmens ab, "
    "durch Differenzierung der direkten Vergleichbarkeit mit Wettbewerbern zu entgehen "
    "(Porter, 2008, S. 93).",
    "Fazit zur Branchenattraktivität:",
)

# -----------------------------------------------------------
# 3.3
# -----------------------------------------------------------
h2("3.3  Wettbewerbsanalyse")

h3("3.3.1  Identifikation der relevanten Wettbewerber")
para(
    "Die Identifikation relevanter Wettbewerber erfordert eine klare Abgrenzung des "
    "Wettbewerbsfelds. Porter (1980, [S. X]) unterscheidet dabei zwischen direkten "
    "Wettbewerbern, die dasselbe Kundensegment mit ähnlichen Produkten und Strategien "
    "bedienen, und indirekten Wettbewerbern, die über substituierende Produkte oder "
    "alternative Lösungsansätze in Konkurrenz treten."
)
para(
    "Für Schaeffler SMB lassen sich folgende Abgrenzungskriterien definieren:"
)
para(
    "Direkte Wettbewerber: Unternehmen, die ebenfalls im Sondermaschinenbau tätig sind, "
    "kundenspezifische Maschinen für vergleichbare Zielindustrien (Automobil, "
    "Präzisionsfertigung) anbieten und eine ähnliche Unternehmensgröße aufweisen. "
    "[Platzhalter – konkrete Wettbewerbernamen auf Basis eigener Marktrecherche ergänzen]",
    "Direkte Wettbewerber:",
)
para(
    "Indirekte Wettbewerber: Anbieter von modularen Automatisierungsplattformen, "
    "Systemintegratoren sowie interne Maschinenbauabteilungen großer Industriekonzerne, "
    "die potenzielle Aufträge im Eigenauftrag realisieren (Porter, 1980, [S. X]).",
    "Indirekte Wettbewerber:",
)

h3("3.3.2  Strategische Gruppen & Positionierung")
para(
    "Das Konzept der strategischen Gruppen beschreibt Cluster von Unternehmen innerhalb "
    "einer Branche, die ähnliche strategische Positionen einnehmen – etwa hinsichtlich "
    "Preisniveau, Qualitätssegment, geografischer Reichweite oder Grad der Spezialisierung "
    "(Porter, 1980, [S. X]). Eine Strategic-Group-Map erlaubt es, die relative "
    "Positionierung von Schaeffler SMB im Wettbewerbsumfeld zu visualisieren."
)
table(
    ["Strategische Gruppe", "Merkmale", "Einordnung Schaeffler SMB"],
    [
        ["Konzernintegrierte\nSondermaschinenbauer",
         "Tochterunternehmen/Abt. großer Industriekonzerne; breites Leistungsspektrum; "
         "Konzernnetzwerk als Vorteil",
         "✓ Zutreffend\n(Schaeffler-Gruppe)"],
        ["Mittelständische\nSpezialisten",
         "Fokus auf wenige Branchen/Technologien; hohes Nischenwissen; begrenzte "
         "Skalierung; projektgetrieben",
         "Teilweise – durch Spezialisierung im Sondermaschinenbau"],
        ["Große Systemintegratoren",
         "Breites Portfolio; Standard- + Sonderlösungen; internationale Präsenz; "
         "hoher Digitalisierungsgrad",
         "Weniger zutreffend – kein breites Standardportfolio"],
        ["Start-ups / Disruptoren",
         "Plattformbasierte Automatisierung; Software-first; häufig ohne eigene Fertigung",
         "Nicht zutreffend"],
    ],
    caption=(
        "Tabelle 3.3: Strategische Gruppen im Sondermaschinenbau. "
        "Quelle: Eigene Darstellung in Anlehnung an Porter (1980, [S. X])."
    ),
)
para(
    "Schaeffler SMB ist aufgrund seiner Konzernanbindung, der globalen Präsenz und des "
    "spezialisierten Fokus primär der ersten strategischen Gruppe zuzuordnen. "
    "[Platzhalter – auf Basis konkreter Wettbewerbsdaten präzisieren; Strategic-Group-Map "
    "als Grafik empfohlen]"
)

h3("3.3.3  Benchmarking-relevante Wettbewerbsprofile")
para(
    "Für den in Kapitel 5 durchgeführten Unternehmensvergleich ist die Auswahl einer "
    "validen Peer-Group entscheidend. Als Benchmarking-Partner kommen Unternehmen in "
    "Frage, die in Bezug auf Branche, Unternehmensgröße, Geschäftsmodell und strategische "
    "Ausrichtung mit Schaeffler SMB vergleichbar sind (Kounev et al., 2020, [S. X])."
)
para("Auswahlkriterien für die Peer-Group:")
bullet("Branchenzugehörigkeit: Sondermaschinenbau oder verwandter Bereich der Fertigungstechnik")
bullet("Unternehmensgröße: Vergleichbare Mitarbeiter- und Umsatzbasis (ca. 500–3.000 MA, Umsatz €50–200 Mio.)")
bullet("Internationaler Footprint: Mindestens europäische, idealerweise globale Präsenz")
bullet("Geschäftsmodell: Engineer-to-Order oder projektbasiertes Maschinengeschäft")
doc.add_paragraph()
table(
    ["Unternehmen", "Sitz", "Größe (MA / Umsatz ca.)", "Auswahlbegründung"],
    [
        ["[Wettbewerber 1]", "[Platzhalter]", "[Platzhalter]", "[Platzhalter]"],
        ["[Wettbewerber 2]", "[Platzhalter]", "[Platzhalter]", "[Platzhalter]"],
        ["[Wettbewerber 3]", "[Platzhalter]", "[Platzhalter]", "[Platzhalter]"],
    ],
    caption="Tabelle 3.4: Benchmarking-Peer-Group Schaeffler SMB. Quelle: Eigene Darstellung.",
)

h3("3.3.4  Wettbewerbsintensität & Differenzierungsmerkmale")
para(
    "Auf Basis der Five-Forces-Analyse (Abschnitt 3.2.4) und der strategischen "
    "Gruppenstruktur lässt sich die Wettbewerbsintensität für Schaeffler SMB differenziert "
    "einschätzen. Porter (1985, [S. X]) betont, dass nachhaltige Wettbewerbsvorteile "
    "entweder durch Kostenführerschaft oder durch Differenzierung entstehen. Im "
    "Sondermaschinenbau ist eine reine Kostenstrategie kaum tragfähig, da die "
    "Individualität des Produkts und die Projektgebundenheit des Geschäfts einer "
    "Massenproduktion entgegenstehen. Differenzierung entsteht daher primär durch "
    "technologische Kompetenz, Kundennähe, Liefergeschwindigkeit und den Grad der "
    "Integration in die Wertschöpfungsprozesse des Kunden (Grant, 2024, [S. X])."
)
para(
    "Schaeffler SMB verfügt durch die Konzernzugehörigkeit über spezifische "
    "Differenzierungsmerkmale: Zugang zu Schaeffler-eigenen Technologien und Komponenten "
    "(insbesondere Wälzlager, Linearsysteme), ein globales Vertriebs- und Servicenetzwerk "
    "sowie das Vertrauen eines etablierten Markennamens. "
    "[Platzhalter – strategische Lücken und konkrete USPs auf Basis interner Analyse ergänzen]"
)

# -----------------------------------------------------------
# 3.4
# -----------------------------------------------------------
h2("3.4  Interne Unternehmensanalyse")
para(
    "Die interne Unternehmensanalyse beleuchtet die Ressourcen, Kompetenzen und "
    "Aktivitäten von Schaeffler SMB, die als Grundlage für Wettbewerbsvorteile dienen. "
    "Sie folgt dem ressourcenbasierten Ansatz sowie der Wertkettenanalyse nach Porter "
    "(1985) und mündet in einer integrierten SWOT-Betrachtung."
)

h3("3.4.1  Ressourcen- & Kompetenzanalyse (Resource-Based View)")
para(
    "Der Resource-Based View (RBV) versteht das Unternehmen als ein Bündel heterogener, "
    "schwer imitierbarer Ressourcen und Kompetenzen, die die Grundlage für nachhaltigen "
    "Wettbewerbsvorteil bilden (Barney, 1991, S. 101). Nach Barney (1991, S. 106) sind "
    "Ressourcen dann strategisch wertvoll, wenn sie die VRIN-Kriterien erfüllen: wertvoll "
    "(valuable), selten (rare), nicht imitierbar (inimitable) und nicht substituierbar "
    "(non-substitutable)."
)
para(
    "Prahalad und Hamel (1990, S. 82) ergänzen diesen Ansatz um das Konzept der "
    "Kernkompetenz: eine kollektiv erworbene organisationale Fähigkeit, die schwer zu "
    "imitieren ist, Zugang zu einer Vielzahl von Märkten schafft und zum wahrgenommenen "
    "Kundenwert beiträgt."
)
table(
    ["Ressourcenkategorie", "Ressourcen Schaeffler SMB", "Strategische Bedeutung (VRIN)"],
    [
        ["Finanziell",
         "Rückhalt der Schaeffler-Gruppe; [Platzhalter: Eigenkapitalquote, Kreditlinie]",
         "Mittel – Absicherung gegen Investitionsschwankungen im Projektgeschäft"],
        ["Physisch",
         "Produktions- und Montagestandorte weltweit; spezialisierte Fertigungsausstattung",
         "Mittel – notwendige Infrastruktur, aber replizierbar"],
        ["Human",
         "Ca. 1.600 Mitarbeitende; hoher Anteil Ingenieure und Facharbeiter; "
         "[Platzhalter: Fluktuationsrate, Ausbildungsquote]",
         "Hoch – kritische Ressource angesichts Fachkräftemangels; schwer replizierbar"],
        ["Immateriell",
         "Schaeffler-Marke; Applikations-Know-how; Patente; langjährige "
         "Kundenbeziehungen; Zertifizierungen",
         "Sehr hoch – VRIN-konform (Barney, 1991); schwer imitierbar"],
    ],
    caption=(
        "Tabelle 3.5: Ressourcenanalyse Schaeffler SMB nach RBV. "
        "Quelle: Eigene Darstellung in Anlehnung an Barney (1991, S. 105–107)."
    ),
)
para(
    "Die Kernkompetenz von Schaeffler SMB liegt in der integrierten Engineering- und "
    "Fertigungskompetenz für hochpräzise Sondermaschinen, kombiniert mit dem Zugang zu "
    "Schaeffler-Komponenten-Know-how. Diese Kombination ist für externe Wettbewerber "
    "schwer zu replizieren (Prahalad & Hamel, 1990, S. 83–84). [Platzhalter – weitere "
    "spezifische Kernkompetenzen aus interner Befragung oder Dokumentenanalyse ergänzen]"
)

h3("3.4.2  Wertkettenanalyse nach Porter")
para(
    "Die Wertkette nach Porter (1985, [S. X]) zergliedert das Unternehmen in seine "
    "wertschöpfenden Aktivitäten und macht sichtbar, wo Kosten entstehen und wo Quellen "
    "von Wettbewerbsvorteil liegen. Sie unterscheidet zwischen primären Aktivitäten, die "
    "direkt zur Leistungserstellung beitragen, und Unterstützungsaktivitäten, die die "
    "Primäraktivitäten ermöglichen."
)
table(
    ["Aktivität", "Beschreibung für Schaeffler SMB", "Kostenrelevanz / Werttreiber"],
    [
        ["PRIMÄRE AKTIVITÄTEN", "", ""],
        ["Eingangslogistik",
         "Beschaffung von Antriebssystemen, Steuerungen, Strukturbauteilen; "
         "Synergien durch Schaeffler-interne Beschaffungsorganisation",
         "Kostenrelevant; Skaleneffekte durch Konzernanbindung"],
        ["Operationen",
         "Konstruktion & Engineering; mechanische Bearbeitung; Montage; "
         "SW-/Steuerungsentwicklung; Qualitätssicherung und Abnahmetests",
         "Höchste Wertentstehung; Kernkompetenz"],
        ["Ausgangslogistik",
         "Transport und Installation beim Kunden weltweit; "
         "Inbetriebnahme und Qualifizierung der Maschinen",
         "Servicequalität als Differenzierungsmerkmal"],
        ["Marketing & Vertrieb",
         "Projektakquise über Ausschreibungen und Direktansprache; "
         "Nutzung des globalen Schaeffler-Vertriebsnetzwerks",
         "Indirekte Wertschöpfung; Konzernsynergie"],
        ["Service",
         "Wartung, Reparatur, Umbau; Ersatzteilversorgung; Remote-Support; "
         "[Platzhalter: Anteil Service am Umsatz]",
         "Wachsender Ertragsbeitrag; Kundenbindung"],
        ["UNTERSTÜTZUNGSAKTIVITÄTEN", "", ""],
        ["Unternehmensinfrastruktur",
         "Management, Controlling, Finanzierung durch Schaeffler-Gruppe; IT-Infrastruktur",
         "Fixkostenblock; Konzernsynergie"],
        ["Personalwirtschaft",
         "Rekrutierung, Weiterbildung, Mitarbeiterbindung – "
         "kritisch in Zeiten des Fachkräftemangels",
         "Strategisch kritisch; Differenzierungspotenzial"],
        ["Technologieentwicklung",
         "F&E; Digitalisierung der Maschinenplattform (IoT, digitaler Zwilling)",
         "Investition in zukünftige Wettbewerbsfähigkeit"],
        ["Beschaffung",
         "Strategisches Lieferantenmanagement; "
         "Nutzung von Skaleneffekten innerhalb der Schaeffler-Gruppe",
         "Kostensenkungspotenzial"],
    ],
    caption=(
        "Tabelle 3.6: Wertkettenanalyse Schaeffler SMB. "
        "Quelle: Eigene Darstellung in Anlehnung an Porter (1985, [S. X])."
    ),
)
para(
    "[Platzhalter – Wertkettenanalyse auf Basis konkreter Kostendaten und Interviews "
    "vertiefen; grafische Darstellung der Wertkette als Abbildung empfohlen]"
)

h3("3.4.3  SWOT-Analyse")
para(
    "Die SWOT-Analyse integriert die Erkenntnisse der externen Umfeldanalyse "
    "(Kapitel 3.2) und der internen Ressourcenanalyse (Kapitel 3.4.1–3.4.2) zu einer "
    "übergreifenden strategischen Beurteilung (Whittington et al., 2024, [S. X]). Sie "
    "identifiziert interne Stärken (Strengths) und Schwächen (Weaknesses) sowie externe "
    "Chancen (Opportunities) und Risiken (Threats)."
)
table(
    ["", "Intern positiv – Stärken (S)", "Intern negativ – Schwächen (W)"],
    [
        ["Extern positiv\nChancen (O)",
         "• Einbindung in globalen Schaeffler-Konzern\n"
         "• Technologische Tiefe im Sondermaschinenbau\n"
         "• Etablierter Kundenstamm (Automobil, Industrie)\n"
         "• Internationales Produktions- und Servicenetz\n"
         "[Platzhalter: weitere Stärken ergänzen]",
         "• [Platzhalter: z. B. Abhängigkeit vom Mutterkonzern]\n"
         "• Begrenzte externe Marketingpräsenz\n"
         "[Platzhalter: weitere Schwächen ergänzen]"],
        ["Extern negativ\nRisiken (T)",
         "• Wachstum durch Automatisierungstrend\n"
         "• Reshoring-Investitionen europäischer Industrie\n"
         "• Nachfrage nach IoT-fähigen Maschinen\n"
         "• Neue Märkte durch Energiewende-Fertigung\n"
         "[Platzhalter: weitere Chancen ergänzen]",
         "• Fachkräftemangel (MINT)\n"
         "• Konjunkturabhängigkeit von Abnehmerindustrien\n"
         "• Geopolitische Risiken (Export, Lieferketten)\n"
         "• Neue Tech-Wettbewerber aus Software/Robotik\n"
         "[Platzhalter: weitere Risiken ergänzen]"],
    ],
    caption="Tabelle 3.7: SWOT-Matrix Schaeffler SMB. Quelle: Eigene Darstellung.",
)

para("Normstrategien-Ableitung:", "Normstrategien-Ableitung:")
table(
    ["", "Chancen (O)", "Risiken (T)"],
    [
        ["Stärken (S)",
         "SO-Strategien:\n"
         "• Schaeffler-Technologie nutzen, um IoT-fähige Maschinen als Wachstumssegment zu erschließen\n"
         "• Globales Schaeffler-Netz für Reshoring-Kunden aktivieren",
         "ST-Strategien:\n"
         "• Schaeffler-Markenstärke als Differenzierungsmerkmal gegen neue Tech-Wettbewerber einsetzen\n"
         "• Langfristige Kundenverträge zur Absicherung gegen Konjunkturschwankungen nutzen"],
        ["Schwächen (W)",
         "WO-Strategien:\n"
         "• Digitalisierungsinvestitionen nutzen, um Effizienzlücken zu schließen und neue datenbasierte Services zu schaffen\n"
         "[Platzhalter: weitere WO-Strategien]",
         "WT-Strategien:\n"
         "• Kundenbranchen diversifizieren zur Risikostreuung\n"
         "• Fachkräftebindung durch Weiterbildungsprogramme stärken"],
    ],
    caption=(
        "Tabelle 3.8: SWOT-Normstrategien Schaeffler SMB. "
        "Quelle: Eigene Darstellung in Anlehnung an Whittington et al. (2024, [S. X])."
    ),
)
para(
    "Die SWOT-Analyse verdeutlicht, dass Schaeffler SMB über eine robuste Ressourcenbasis "
    "verfügt, die durch die Konzerneinbindung gestärkt wird. Die größten strategischen "
    "Herausforderungen liegen in der Bewältigung des Fachkräftemangels, der "
    "Konjunkturabhängigkeit und den notwendigen Digitalisierungsinvestitionen. Die "
    "identifizierten SO-Strategien bilden die Grundlage für die Benchmarking-Kennzahlen "
    "in Kapitel 5."
)

# -----------------------------------------------------------
# LITERATURVERZEICHNIS
# -----------------------------------------------------------
doc.add_page_break()
h1("Literaturverzeichnis (Kapitel 3)")

bib = [
    ("Barney, J. (1991). Firm resources and sustained competitive advantage. ",
     "Journal of Management",
     ", 17(1), 99–120. https://doi.org/10.1177/014920639101700108"),
    ("Grant, R. M. (2024). ",
     "Contemporary strategy analysis",
     " (12. Aufl.). Wiley. ISBN 978-1-394-25159-9"),
    ("Kounev, S., Lange, K.-D., & von Kistowski, J. (2020). ",
     "Systems benchmarking: For scientists and engineers.",
     " Springer International Publishing. https://doi.org/10.1007/978-3-030-41705-5"),
    ("Porter, M. E. (1980). ",
     "Competitive strategy: Techniques for analyzing industries and competitors.",
     " The Free Press. ISBN 978-0-02-925360-1"),
    ("Porter, M. E. (1985). ",
     "Competitive advantage: Creating and sustaining superior performance.",
     " The Free Press. ISBN 978-0-02-925090-7"),
    ("Porter, M. E. (2008). The five competitive forces that shape strategy. ",
     "Harvard Business Review",
     ", 86(1), 78–93."),
    ("Prahalad, C. K., & Hamel, G. (1990). The core competence of the corporation. ",
     "Harvard Business Review",
     ", 68(3), 79–91."),
    ("VDMA. (2025). ",
     "Jahresbilanz Maschinenbau 2024.",
     " Verband Deutscher Maschinen- und Anlagenbau. https://www.vdma.org"),
    ("Whittington, R., Regnér, P., Angwin, D., Johnson, G., & Scholes, K. (2024). ",
     "Exploring strategy: Text and cases",
     " (13. Aufl.). Pearson Education. ISBN 978-1-292-47953-8"),
]

for before, italic_part, after in bib:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent   = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    p.paragraph_format.space_after   = Pt(4)
    r1 = p.add_run(before)
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    r2 = p.add_run(italic_part)
    r2.font.name   = 'Times New Roman'
    r2.font.size   = Pt(12)
    r2.font.italic = True
    r3 = p.add_run(after)
    r3.font.name = 'Times New Roman'
    r3.font.size = Pt(12)

# -----------------------------------------------------------
# SPEICHERN
# -----------------------------------------------------------
out = '/home/user/test-data1/Kapitel_3_Wettbewerbsanalyse.docx'
doc.save(out)
print(f"Dokument gespeichert: {out}")
