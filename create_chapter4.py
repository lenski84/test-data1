from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# Seitenränder DIN A4
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2)

FONT = "Times New Roman"
BASE = 12

def set_font(run, size=BASE, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), FONT)
    rFonts.set(qn('w:hAnsi'), FONT)
    rFonts.set(qn('w:cs'),    FONT)
    rPr.insert(0, rFonts)

def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18) if level == 1 else Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run(text)
    size = 14 if level == 1 else 12
    set_font(run, size=size, bold=True)
    return p

def body(doc, text, space_after=6, justify=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(18)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run)
    return p

def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run("─" * 80)
    run.font.name = FONT
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

def add_table_row(table, cells_data, bold=False, bg=None):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        set_font(run, bold=bold)
        if bg:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), bg)
            tcPr.append(shd)
    return row

def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[i])

# ─────────────────────────────────────────────
# KAPITEL 4: METHODISCHES VORGEHEN
# ─────────────────────────────────────────────

heading(doc, "4  Methodisches Vorgehen", level=1)

body(doc,
    "Das vorliegende Kapitel beschreibt das methodische Vorgehen der Untersuchung. "
    "Es umfasst die Begründung des Forschungsdesigns, die Auswahl des zugrunde liegenden "
    "Benchmarking-Modells, die Kriterien zur Auswahl der Vergleichsunternehmen, "
    "die eingesetzten Methoden der Datenerhebung und -auswertung sowie eine kritische "
    "Reflexion der Grenzen der Untersuchung.")

divider(doc)

# ─── 4.1 Forschungsdesign ───
heading(doc, "4.1  Forschungsdesign", level=2)

body(doc,
    "Die Arbeit folgt einem angewandten, deskriptiv-analytischen Forschungsdesign. "
    "Ziel ist nicht die Generierung neuer Theorie, sondern die strukturierte Beschreibung "
    "und Bewertung eines Ist-Zustands sowie die Ableitung konkreter Handlungsempfehlungen "
    "für ein Unternehmen der Praxis (Schaeffler Special Machinery). Damit ordnet sich "
    "die Untersuchung in den Bereich der angewandten Wirtschaftsforschung ein "
    "(Bortz & Döring, 2016, S. 48).")

body(doc,
    "Als Forschungsstrategie wird das Einzelfallstudiendesign (Single Case Study) "
    "nach Yin (2018) gewählt. Dieses eignet sich dann, wenn ein spezifisches Phänomen "
    "in seinem realen Unternehmenskontext untersucht wird und eine Generalisierbarkeit "
    "der Ergebnisse nicht primäres Ziel ist (Yin, 2018, S. 29). "
    "Schaeffler SMB stellt hier den fokalen Fall dar; die Wettbewerber dienen als "
    "Vergleichseinheiten im Rahmen des externen Benchmarks.")

body(doc,
    "Das Vorgehen ist primär deduktiv: Die in Kapitel 2 und 3 erarbeiteten theoretischen "
    "Grundlagen – insbesondere das Fraunhofer-IZB-Fünf-Phasen-Modell sowie die "
    "strategischen Analyseinstrumente – bilden den konzeptionellen Rahmen, "
    "in den die erhobenen empirischen Daten eingeordnet werden.")

divider(doc)

# ─── 4.2 Benchmarking-Modell ───
heading(doc, "4.2  Auswahl des Benchmarking-Modells", level=2)

body(doc,
    "Für die strukturierte Durchführung des Benchmarks wird das Fünf-Phasen-Modell "
    "des Fraunhofer-Instituts für Zuverlässigkeit und Mikrointegration (IZB) als "
    "Leitrahmen verwendet. Dieses Modell hat sich im deutschsprachigen Industrieumfeld "
    "als Standard etabliert und ist speziell auf produzierende Unternehmen ausgelegt "
    "(Fraunhofer IZB, 2023). Es ist strukturell mit dem Zehn-Schritte-Modell von "
    "Camp (1994) kompatibel, dem Ursprungsmodell des modernen Benchmarkings aus "
    "dem Xerox-Kontext.")

body(doc,
    "Die fünf Phasen umfassen: (1) Initialisierung und Zieldefinition, "
    "(2) interne Analyse, (3) Vergleichspartnerauswahl und Datenerhebung, "
    "(4) Vergleich und Gap-Analyse sowie (5) Maßnahmenplanung und Umsetzung "
    "(Fraunhofer IZB, 2023, S. 14–17). "
    "Die vorliegende Projektarbeit deckt die Phasen 1 bis 4 vollständig ab; "
    "Phase 5 wird im Rahmen der strategischen Handlungsempfehlungen (Kapitel 7) "
    "konzeptionell adressiert, ohne operative Umsetzungsverantwortung zu beanspruchen.")

# Tabelle: Phasen → Kapitel-Mapping
body(doc, "Tabelle 4.1: Zuordnung der Fraunhofer-IZB-Phasen zu den Kapiteln der Arbeit",
     space_after=3)

tbl = doc.add_table(rows=1, cols=3)
tbl.style = "Table Grid"
add_table_row(tbl, ["Phase", "Inhalt", "Kapitel"], bold=True, bg="D9D9D9")
data = [
    ("1", "Initialisierung & Zieldefinition", "Kap. 1, 4"),
    ("2", "Interne Analyse (Selbstbewertung)", "Kap. 5"),
    ("3", "Vergleichspartnerauswahl & Datenerhebung", "Kap. 4, 5"),
    ("4", "Vergleich & Gap-Analyse", "Kap. 6"),
    ("5", "Maßnahmenplanung & Umsetzung", "Kap. 7"),
]
for row in data:
    add_table_row(tbl, row)
set_col_widths(tbl, [1.5, 9.0, 3.0])
doc.add_paragraph()

divider(doc)

# ─── 4.3 Vergleichspartnerauswahl ───
heading(doc, "4.3  Auswahl der Vergleichsunternehmen", level=2)

body(doc,
    "Die Auswahl geeigneter Vergleichsunternehmen ist eine der zentralen methodischen "
    "Herausforderungen des externen Benchmarkings, insbesondere im Segment des "
    "Sondermaschinenbaus. Da Sondermaschinenbauer typischerweise Einzelfertigungen "
    "oder Kleinserienfertiger sind, fehlen standardisierte Leistungskennzahlen, "
    "und eine direkte Vergleichbarkeit ist oft nur eingeschränkt gegeben "
    "(Spendolini, 1992, S. 78).")

body(doc,
    "Für die vorliegende Arbeit wurden Vergleichsunternehmen anhand folgender "
    "Auswahlkriterien identifiziert:")

# Aufzählung
for item in [
    "Branchenzugehörigkeit: Maschinenbau / Sondermaschinenbau (NACE-Code 28.x)",
    "Umsatzgröße: vergleichbare Größenordnung (mittlerer bis oberer Mittelstand)",
    "Geschäftsmodell: Auftragsfertiger mit kundenspezifischen Lösungen",
    "Internationales Tätigkeitsprofil: Exportquote und globale Marktpräsenz",
    "Datenverfügbarkeit: Jahresberichte, Unternehmensregister, Fachpublikationen",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(1)
    run = p.add_run(item)
    set_font(run)

doc.add_paragraph()

body(doc,
    "Das Problem der Vergleichbarkeit (Komparabilität) wird durch ein zweigliedriges "
    "Vorgehen adressiert: Erstens werden absolute Kennzahlen (z. B. Umsatz, Mitarbeiterzahl) "
    "stets durch relative Kennzahlen (z. B. Umsatz je Mitarbeiter, F&E-Quote) ergänzt, "
    "um Größeneffekte zu neutralisieren. Zweitens wird die Auswahl der Vergleichsdimensionen "
    "auf strategisch relevante Parameter beschränkt, die branchenübergreifend interpretierbar "
    "sind (Watson, 1993, S. 61).")

body(doc,
    "In Übereinstimmung mit Camp (1994, S. 67) wurde zudem ein sog. "
    "Best-in-Class-Vergleich angestrebt: Neben direkten Wettbewerbern wurden, "
    "wo methodisch sinnvoll, auch branchenfremde Unternehmen mit übertragbaren "
    "Prozessexzellenzen berücksichtigt (funktionales Benchmarking). "
    "Die konkrete Auswahl der Vergleichsunternehmen sowie deren Steckbriefe "
    "sind in Kapitel 5 dokumentiert.")

divider(doc)

# ─── 4.4 Datenerhebung ───
heading(doc, "4.4  Datenerhebung", level=2)

body(doc,
    "Die Datenerhebung erfolgt auf Basis einer kombinierten Sekundär- und "
    "Primärforschungsstrategie. Aufgrund des vertraulichen Charakters der Untersuchung "
    "und der eingeschränkten Verfügbarkeit öffentlicher Daten im Sondermaschinenbau "
    "bildet die Sekundärforschung den Schwerpunkt.")

heading(doc, "4.4.1  Sekundärquellen", level=2)

body(doc,
    "Als Sekundärquellen wurden systematisch folgende Quellentypen ausgewertet:")

# Tabelle Sekundärquellen
tbl2 = doc.add_table(rows=1, cols=3)
tbl2.style = "Table Grid"
add_table_row(tbl2, ["Quellentyp", "Beispiele", "Bewertung"], bold=True, bg="D9D9D9")
src_data = [
    ("Jahres-/Geschäftsberichte",
     "Veröffentlichte Jahresabschlüsse der Vergleichsunternehmen",
     "Hoch verlässlich, aber eingeschränkter Detailgrad"),
    ("Branchenreports & Verbandsdaten",
     "VDMA Maschinenbau in Zahl und Bild 2024",
     "Repräsentativ, aktuell, institutionell validiert"),
    ("Unternehmensregister",
     "Bundesanzeiger, Creditreform, Dun & Bradstreet",
     "Standardisiert, vergleichbar"),
    ("Fachliteratur & Studien",
     "Fraunhofer-Publikationen, IHK-Studien",
     "Methodisch fundiert"),
    ("Unternehmenswebsites & PR",
     "Produktportfolios, Pressemitteilungen",
     "Ergänzend, kritisch zu prüfen"),
]
for row in src_data:
    add_table_row(tbl2, row)
set_col_widths(tbl2, [4.0, 6.0, 5.0])
doc.add_paragraph()

body(doc,
    "Branchenkennzahlen des deutschen Maschinenbaus wurden primaer aus der Publikation "
    "Maschinenbau in Zahl und Bild 2024 des Verbandes Deutscher Maschinen- und "
    "Anlagenbau (VDMA) entnommen. Der VDMA erhebt diese Kennzahlen auf Basis einer "
    "repraesentativen Mitgliederbefragung und gilt als institutionell verlaesslichste "
    "Quelle fuer diese Branche in Deutschland (VDMA, 2024).")

heading(doc, "4.4.2  Primärquellen", level=2)

body(doc,
    "Ergänzend zur Sekundärforschung wurden unternehmensinterne Primärdaten "
    "aus dem Bereich Strategy & Processes der Schaeffler Special Machinery einbezogen. "
    "Diese umfassen Kennzahlen aus dem internen Berichtswesen sowie qualitative "
    "Einschätzungen aus Expertengesprächen mit Führungskräften der Abteilung. "
    "Die Primärdaten unterliegen dem Sperrvermerk dieser Arbeit und sind daher "
    "in aggregierter, anonymisierter Form dargestellt.")

body(doc,
    "Qualitative Primärdaten wurden im Rahmen von leitfadengestützten Experteninterviews "
    "erhoben. Die Auswahl der Gesprächspartner folgte dem Prinzip des Purposive Sampling: "
    "Es wurden Personen befragt, die aufgrund ihrer Funktion über fundiertes Wissen "
    "zum Untersuchungsgegenstand verfügen (Yin, 2018, S. 113). "
    "Die Interviews wurden strukturiert protokolliert und inhaltsanalytisch ausgewertet.")

divider(doc)

# ─── 4.5 Datenauswertung ───
heading(doc, "4.5  Datenauswertung und Kennzahlensystem", level=2)

body(doc,
    "Die Auswertung der erhobenen Daten erfolgt anhand eines strukturierten "
    "Kennzahlensystems. Dieses umfasst quantitative Leistungsindikatoren (Key Performance "
    "Indicators, KPIs) sowie qualitative Bewertungsdimensionen. Die Kombination beider "
    "Ebenen entspricht dem Ansatz eines hybriden Benchmarks, der sowohl messbare "
    "Größen als auch strategische Qualitätsdimensionen berücksichtigt "
    "(Camp, 1994, S. 102).")

body(doc,
    "Das Kennzahlensystem ist entlang der vier strategischen Vergleichsdimensionen "
    "gegliedert, die aus der Zielsetzung der Arbeit abgeleitet wurden:")

# Kennzahlensystem-Tabelle
tbl3 = doc.add_table(rows=1, cols=3)
tbl3.style = "Table Grid"
add_table_row(tbl3, ["Dimension", "Beispiel-KPIs", "Datenquelle"], bold=True, bg="D9D9D9")
kpi_data = [
    ("Markt & Wachstum",
     "Umsatzwachstum (%), Marktanteil, Exportquote",
     "Jahresbericht, VDMA"),
    ("Kosten & Effizienz",
     "Umsatz/Mitarbeiter, EBIT-Marge, Materialquote",
     "Jahresbericht, Creditreform"),
    ("Leistung & Portfolio",
     "Produktbreite, Technologiereife, Zertifizierungen",
     "Websites, Experteninterviews"),
    ("Strategie & Innovation",
     "F&E-Quote, Patentanzahl, Digitalisierungsgrad",
     "Geschäftsberichte, Fachpresse"),
]
for row in kpi_data:
    add_table_row(tbl3, row)
set_col_widths(tbl3, [4.0, 6.5, 4.5])
doc.add_paragraph()

body(doc,
    "Die Gap-Analyse vergleicht die Ist-Werte von Schaeffler SMB mit den "
    "ermittelten Benchmarkwerten der Vergleichsunternehmen. Gaps werden sowohl "
    "absolut als auch prozentual ausgewiesen. Zur Priorisierung der identifizierten "
    "Lücken wird eine Relevanzbewertung anhand der Kriterien strategische Bedeutung "
    "und Schließbarkeit vorgenommen (Watson, 1993, S. 84).")

divider(doc)

# ─── 4.6 Grenzen ───
heading(doc, "4.6  Grenzen der Untersuchung", level=2)

body(doc,
    "Die Methodik der vorliegenden Arbeit unterliegt verschiedenen Einschränkungen, "
    "die bei der Interpretation der Ergebnisse zu berücksichtigen sind.")

body(doc,
    "Erstens ist die Datenverfügbarkeit im Sondermaschinenbau strukturell begrenzt. "
    "Viele Unternehmen dieses Segments sind nicht börsennotiert und unterliegen "
    "keiner umfassenden Publizitätspflicht. Detaillierte operative Kennzahlen "
    "(z. B. Fertigungstiefe, Projektlaufzeiten) sind daher häufig nicht öffentlich "
    "zugänglich und mussten durch Schätzwerte oder Branchendurchschnitte approximiert "
    "werden (Spendolini, 1992, S. 79).")

body(doc,
    "Zweitens besteht eine inhärente Vergleichbarkeitsproblematik: Selbst Unternehmen "
    "derselben Branche unterscheiden sich hinsichtlich Fertigungstiefe, Kundenstruktur, "
    "regionaler Ausrichtung und Produktkomplexität so stark, dass ein direkter "
    "Kennzahlenvergleich stets relativiert werden muss. Die in Abschnitt 4.3 "
    "beschriebenen Normierungsmaßnahmen mildern dieses Problem, eliminieren es "
    "jedoch nicht vollständig.")

body(doc,
    "Drittens ist der Erhebungszeitraum auf das Geschäftsjahr 2023/2024 beschränkt. "
    "Angesichts der hohen Dynamik im deutschen Maschinenbau – geprägt durch "
    "geopolitische Verwerfungen, Energiepreisschocks und Nachfrageeinbrüche "
    "aus China – können die Ergebnisse eine eingeschränkte zeitliche Gültigkeit "
    "aufweisen (VDMA, 2024, S. 5).")

body(doc,
    "Viertens basieren qualitative Einschätzungen teilweise auf Expertenmeinungen "
    "innerhalb von Schaeffler SMB, die naturgemäß eine unternehmensinterne Perspektive "
    "widerspiegeln. Eine vollständig externe Validierung aller qualitativen "
    "Bewertungen war im Rahmen dieser Projektarbeit nicht möglich.")

divider(doc)

# ─── LITERATURVERZEICHNIS ───
heading(doc, "Literaturverzeichnis (Kapitel 4)", level=1)

refs = [
    ("Bortz, J., & Döring, N. (2016). "
     "Forschungsmethoden und Evaluation für Human- und Sozialwissenschaftler "
     "(5., überarb. Aufl.). Springer."),
    ("Camp, R. C. (1994). "
     "Benchmarking: The Search for Industry Best Practices that Lead to Superior Performance. "
     "ASQC Quality Press."),
    ("Fraunhofer IZB. (2023). "
     "Benchmarking im Mittelstand: Das Fraunhofer-Modell. "
     "Fraunhofer-Institut für Zuverlässigkeit und Mikrointegration."),
    ("Spendolini, M. J. (1992). "
     "The Benchmarking Book. "
     "AMACOM."),
    ("VDMA. (2024). "
     "Maschinenbau in Zahl und Bild 2024. "
     "Verband Deutscher Maschinen- und Anlagenbau e. V."),
    ("Watson, G. H. (1993). "
     "Strategic Benchmarking: How to Rate Your Company's Performance Against the World's Best. "
     "Wiley."),
    ("Yin, R. K. (2018). "
     "Case Study Research and Applications: Design and Methods (6th ed.). "
     "SAGE Publications."),
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent   = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-1.25)
    p.paragraph_format.space_after   = Pt(4)
    p.paragraph_format.line_spacing  = Pt(18)
    run = p.add_run(ref)
    set_font(run, italic=False)

doc.save("Kapitel_4_Methodisches_Vorgehen.docx")
print("Gespeichert: Kapitel_4_Methodisches_Vorgehen.docx")
