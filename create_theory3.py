from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

sec = doc.sections[0]
sec.page_height = Cm(29.7)
sec.page_width  = Cm(21.0)
sec.left_margin   = Cm(3.0)
sec.right_margin  = Cm(2.5)
sec.top_margin    = Cm(2.5)
sec.bottom_margin = Cm(2.0)

normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)
normal.paragraph_format.space_after  = Pt(6)
normal.paragraph_format.line_spacing = Pt(18)

# ── helpers ──────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.name = 'Times New Roman'; r.font.size = Pt(14)
        r.font.bold = True; r.font.color.rgb = RGBColor(0,0,0)
    p.paragraph_format.space_before = Pt(20); p.paragraph_format.space_after = Pt(8)

def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.name = 'Times New Roman'; r.font.size = Pt(13)
        r.font.bold = True; r.font.color.rgb = RGBColor(0,0,0)
    p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(6)

def h3(text):
    p = doc.add_heading(text, level=3)
    for r in p.runs:
        r.font.name = 'Times New Roman'; r.font.size = Pt(12)
        r.font.bold = True; r.font.color.rgb = RGBColor(0,0,0)
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)

def para(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix); r.bold = True
        r.font.name = 'Times New Roman'; r.font.size = Pt(12)
        r2 = p.add_run(text[len(bold_prefix):])
        r2.font.name = 'Times New Roman'; r2.font.size = Pt(12)
    else:
        r = p.add_run(text)
        r.font.name = 'Times New Roman'; r.font.size = Pt(12)

def caption(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(10)
    for r in p.runs:
        r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.font.italic = True

def tbl(headers, rows, cap=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.paragraphs[0].clear()
        r = c.paragraphs[0].add_run(h); r.bold = True
        r.font.name = 'Times New Roman'; r.font.size = Pt(10)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        for ci, txt in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.paragraphs[0].clear()
            r = c.paragraphs[0].add_run(txt)
            r.font.name = 'Times New Roman'; r.font.size = Pt(10)
    if cap:
        caption(cap)
    else:
        doc.add_paragraph()

def divider():
    p = doc.add_paragraph('─' * 80)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(14)
    for r in p.runs:
        r.font.name = 'Times New Roman'; r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(150,150,150)

def bib_entry(before, italic, after):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent        = Cm(1.0)
    p.paragraph_format.first_line_indent  = Cm(-1.0)
    p.paragraph_format.space_after        = Pt(4)
    for txt, itl in [(before, False),(italic, True),(after, False)]:
        r = p.add_run(txt)
        r.font.name = 'Times New Roman'; r.font.size = Pt(12)
        r.font.italic = itl

# ═══════════════════════════════════════════════════════════════════════════
# TITEL
# ═══════════════════════════════════════════════════════════════════════════
h1("Theoretische Grundlagen der Analysemethoden – Kapitel 3")
para(
    "Das vorliegende Kapitel legt die methodisch-theoretische Basis für die in "
    "Kapitel 3 eingesetzten Analyseinstrumente. Für jedes Instrument werden Herkunft, "
    "konzeptionelle Grundlage und Auswertungslogik erläutert, um die spätere Anwendung "
    "auf das Unternehmen und seine Branche intersubjektiv nachvollziehbar zu machen."
)

# ═══════════════════════════════════════════════════════════════════════════
# PESTEL
# ═══════════════════════════════════════════════════════════════════════════
divider()
h2("PESTEL-Analyse")

para(
    "Die PESTEL-Analyse ist ein Framework zur strukturierten Erfassung der "
    "makroökonomischen Umwelt, in der ein Unternehmen oder eine Branche operiert. "
    "Ihr Ursprung liegt in der Arbeit von Francis Aguilar, der 1967 in "
    "„Scanning the Business Environment“ erstmals eine systematische "
    "Kategorisierung externer Einflussfaktoren vorschlug (ETPS: Economic, Technical, "
    "Political, Social). Das Akronym wurde in der Folge zu PEST umgestellt und "
    "schrittweise um die Kategorien Environmental (Ö) und Legal (L) ergänzt, um "
    "wachsenden ökologischen Regulierungsdruck und rechtliche Rahmenbedingungen "
    "eigenständig abzubilden (Whittington et al., 2024, [S. X])."
)
para(
    "Der konzeptionelle Kern des Instruments liegt in der Unterscheidung zwischen "
    "der Mikroumwelt – dem direkt beeinflussbaren Wettbewerbs- und Marktumfeld – "
    "und der Makroumwelt, auf die das Unternehmen keinen unmittelbaren Einfluss hat, "
    "deren Kräfte jedoch seine strategischen Handlungsspielräume grundlegend "
    "bestimmen. PESTEL operiert ausschließlich auf der Makroebene und liefert damit "
    "den Kontext, in dem Branchenanalysen wie die Five Forces oder "
    "unternehmensinterne Analysen erst sinnvoll interpretiert werden können "
    "(Whittington et al., 2024, [S. X])."
)

h3("Aufbau und Auswertungslogik")
para(
    "Die Analyse folgt keinem starren Schema, sondern einer inhaltlichen Prüflogik: "
    "Für jede der sechs Dimensionen werden diejenigen Faktoren identifiziert, die "
    "für das konkrete Unternehmen oder die Branche tatsächlich entscheidungsrelevant "
    "sind. Whittington et al. (2024, [S. X]) warnen ausdrücklich vor einer "
    "vollständigen Enumeration aller denkbaren Faktoren, da dies die strategische "
    "Aussagekraft des Instruments verwässert. Stattdessen sollte jeder identifizierte "
    "Faktor hinsichtlich seiner Relevanz und seiner absehbaren Entwicklungsrichtung "
    "bewertet werden."
)
tbl(
    ["Dimension", "Gegenstand", "Typische Analysefragen"],
    [
        ["Politisch (P)",
         "Staatliche und überstaatliche Eingriffe, Handelspolitik, Subventionen",
         "Welche Regulierungen oder politischen Entscheidungen beeinflussen die Branche? Gibt es Exportbeschränkungen?"],
        ["Ökonomisch (E)",
         "Konjunktur, Zinsen, Wechselkurse, Investitionsklima",
         "In welchem Konjunkturzyklus befindet sich die Branche? Wie reagiert die Nachfrage auf Zinsschwankungen?"],
        ["Sozio-kulturell (S)",
         "Demografischer Wandel, Wertewandel, Arbeitskräfteangebot",
         "Verändert sich die Verfügbarkeit qualifizierter Arbeitskräfte? Wie wandeln sich Kundenpräferenzen?"],
        ["Technologisch (T)",
         "Innovationsgeschwindigkeit, Digitalisierung, neue Produktionstechnologien",
         "Welche Technologien könnten bestehende Geschäftsmodelle verdrängen oder neue ermöglichen?"],
        ["Ökologisch (E)",
         "Klimaregulierung, Ressourcenverfügbarkeit, Nachhaltigkeitsanforderungen",
         "Welche Umweltauflagen gelten? Wie verändert sich der regulatorische Druck auf Energieverbrauch?"],
        ["Legal (L)",
         "Gesetzgebung, Normen, Haftungsrecht, Compliance",
         "Welche Gesetze und Normen sind für Produkt, Vertrieb und Betrieb relevant?"],
    ],
    cap="Tabelle T.1: PESTEL-Dimensionen und Analysefragen. Quelle: Eigene Darstellung in Anlehnung an Whittington et al. (2024, [S. X]).",
)
para(
    "Die Ergebnisse der PESTEL-Analyse werden üblicherweise als externe Eingangsgröße "
    "in die SWOT-Analyse überführt: Makrofaktoren mit positivem Einfluss auf das "
    "Unternehmen gehen als Chancen (O), solche mit negativem Einfluss als Risiken (T) "
    "in die SWOT-Matrix ein."
)

# ═══════════════════════════════════════════════════════════════════════════
# FIVE FORCES
# ═══════════════════════════════════════════════════════════════════════════
divider()
h2("Branchenstrukturanalyse nach Porter – Five Forces")

para(
    "Michael E. Porter entwickelte das Fünf-Kräfte-Modell 1980 in "
    "„Competitive Strategy: Techniques for Analyzing Industries and "
    "Competitors“. Theoretisch fußt es auf dem Structure-Conduct-Performance-"
    "Paradigma der Harvard Industrial Organization: Die Struktur einer Branche "
    "bestimmt das Verhalten der Unternehmen darin, das wiederum deren Performance "
    "determiniert (Porter, 1980, [S. X]). Porter übertrug diesen industrieökonomischen "
    "Ansatz auf das strategische Management und machte ihn damit zum zentralen "
    "Instrument der Branchenanalyse. 2008 aktualisierte er das Modell in einem "
    "Grundlagenbeitrag im Harvard Business Review (Porter, 2008, S. 78)."
)
para(
    "Die Kernthese lautet: Die langfristige Rentabilität einer Branche wird nicht "
    "durch kurzfristige Marktentwicklungen bestimmt, sondern durch fünf strukturelle "
    "Kräfte, die gemeinsam den Wettbewerbsdruck und damit das Gewinnpotenzial "
    "der Branchenanbieter determinieren (Porter, 2008, S. 79). Das Modell beantwortet "
    "damit die Frage, warum manche Branchen dauerhaft profitabler sind als andere – "
    "unabhängig von der Leistung einzelner Unternehmen."
)

h3("Die fünf Kräfte")
tbl(
    ["Kraft", "Definition", "Einflussfaktoren"],
    [
        ["Rivalität der\nbestehenden\nWettbewerber",
         "Intensität des direkten Wettbewerbs unter den etablierten Anbietern",
         "Anzahl und Größe der Wettbewerber, Marktwachstum, Produktdifferenzierung, Wechselkosten, Austrittsbarrieren"],
        ["Verhandlungsmacht\nder Lieferanten",
         "Fähigkeit der Lieferanten, Preise zu erhöhen oder Qualität zu senken",
         "Konzentration der Lieferanten, Verfügbarkeit von Alternativen, Bedeutung des Abnehmers für den Lieferanten, Wechselkosten"],
        ["Verhandlungsmacht\nder Kunden",
         "Fähigkeit der Kunden, Preise zu drücken oder Mehrleistungen zu fordern",
         "Konzentration der Kunden, Volumen, Produktstandardisierung, Wechselkosten, Rückwärtsintegrationspotenzial"],
        ["Bedrohung durch\nSubstitute",
         "Gefahr, dass alternative Produkte oder Lösungen die Branchenleistung ersetzen",
         "Preis-Leistungs-Verhältnis der Substitute, Wechselbereitschaft der Kunden, Technologiedynamik"],
        ["Bedrohung durch\nNeueintreter",
         "Gefahr, dass neue Anbieter in die Branche eintreten und Marktanteile gewinnen",
         "Eintrittsbarrieren (Kapital, Know-how, Regulierung, Skaleneffekte, Markenidentität), Reaktion der Etablierten"],
    ],
    cap="Tabelle T.2: Die fünf Wettbewerbskräfte nach Porter. Quelle: Eigene Darstellung in Anlehnung an Porter (1980, [S. X]); Porter (2008, S. 79–86).",
)

h3("Auswertungslogik")
para(
    "Jede der fünf Kräfte wird anhand branchenspezifischer Indikatoren auf einer "
    "qualitativen Skala (niedrig – mittel – hoch) bewertet. Abschließend werden die "
    "Einzelbewertungen zu einem Gesamturteil über die Branchenattraktivität "
    "synthetisiert: Je stärker die Kräfte insgesamt wirken, desto geringer ist das "
    "strukturelle Gewinnpotenzial der Branche (Porter, 2008, S. 93). Porter (1980, "
    "[S. X]) betont, dass die Analyse auf die Branchenebene und nicht auf das "
    "Einzelunternehmen zielt – sie ist damit Voraussetzung, aber nicht Ersatz für "
    "die Wettbewerbsanalyse auf Unternehmensebene."
)

# ═══════════════════════════════════════════════════════════════════════════
# STRATEGISCHE GRUPPEN
# ═══════════════════════════════════════════════════════════════════════════
divider()
h2("Konzept der Strategischen Gruppen")

para(
    "Das Konzept der strategischen Gruppen wurde von Michael E. Porter (1980) in "
    "„Competitive Strategy“ eingeführt und basiert auf frühen "
    "industrieökonomischen Arbeiten von Hunt (1972), der Strukturheterogenität "
    "innerhalb von Branchen empirisch beobachtet hatte. Porter stellte fest, dass "
    "Unternehmen innerhalb einer Branche selten homogen konkurrieren, sondern sich "
    "in Clustern ähnlicher strategischer Ausrichtung zusammenfinden. Diese Cluster "
    "bezeichnet er als strategische Gruppen (Porter, 1980, [S. X])."
)
para(
    "Eine strategische Gruppe ist definiert als eine Menge von Unternehmen, die "
    "innerhalb einer Branche dieselbe oder eine ähnliche Strategie entlang "
    "strategischer Dimensionen verfolgen (Porter, 1980, [S. X]). Der analytische "
    "Nutzen liegt darin, dass der Wettbewerb innerhalb einer strategischen Gruppe "
    "in der Regel intensiver ist als zwischen verschiedenen Gruppen, da die "
    "Unternehmen um dieselben Kunden mit ähnlichen Mitteln konkurrieren. "
    "Gleichzeitig werden durch die Kartierung strategische Lücken – unbesetzte "
    "Positionen im Wettbewerbsraum – sichtbar (Grant, 2024, [S. X])."
)

h3("Strategic-Group-Map: Methodik")
para(
    "In einer Strategic-Group-Map werden Unternehmen anhand zweier strategisch "
    "relevanter Dimensionen in einem zweidimensionalen Raum verortet. Die Auswahl "
    "der Achsen ist entscheidend für die Aussagekraft der Karte: Sie sollte auf "
    "denjenigen Merkmalen basieren, die in der jeweiligen Branche die bedeutsamsten "
    "Differenzierungslinien darstellen – etwa Preisniveau und Produktbreite, "
    "geografische Reichweite und Spezialisierungsgrad oder Fertigungstiefe und "
    "Markenorientierung (Porter, 1980, [S. X]). Unternehmen, die in beiden "
    "Dimensionen ähnliche Positionen einnehmen, werden zu einer strategischen "
    "Gruppe zusammengefasst, die üblicherweise durch einen Kreis visualisiert wird, "
    "dessen Größe den relativen Marktanteil der Gruppe widerspiegelt."
)

# ═══════════════════════════════════════════════════════════════════════════
# GENERISCHE WETTBEWERBSSTRATEGIEN
# ═══════════════════════════════════════════════════════════════════════════
divider()
h2("Generische Wettbewerbsstrategien und Differenzierung nach Porter")

para(
    "Aufbauend auf der Branchenstrukturanalyse entwickelte Porter (1985) in "
    "„Competitive Advantage“ das Konzept der generischen Wettbewerbs-"
    "strategien. Es beantwortet die Frage, wie ein einzelnes Unternehmen innerhalb "
    "einer gegebenen Branchenstruktur eine überlegene Wettbewerbsposition "
    "aufbauen und verteidigen kann (Porter, 1985, [S. X])."
)
para(
    "Porter unterscheidet grundlegend zwei Quellen von Wettbewerbsvorteilen: "
    "niedrigere Kosten als die Konkurrenz (Kostenführerschaft) oder einen "
    "wahrnehmbaren Mehrwert für den Kunden, für den dieser bereit ist, einen "
    "Aufpreis zu zahlen (Differenzierung). In Kombination mit dem angestrebten "
    "Marktumfang – breiter Massenmarkt oder fokussiertes Nischensegment – "
    "ergeben sich drei generische Strategien: Kostenführerschaft, Differenzierung "
    "und Fokus (Porter, 1985, [S. X])."
)
tbl(
    ["Strategie", "Grundprinzip", "Voraussetzungen"],
    [
        ["Kostenführerschaft",
         "Günstigster Anbieter in der Branche durch Skaleneffekte, Prozessoptimierung oder Einkaufsmacht",
         "Hohe Automatisierung, Standardisierung, Kostenkontrolle, hohe Absatzmengen"],
        ["Differenzierung",
         "Einzigartiges Leistungsangebot, für das Kunden einen Preis-Aufschlag zahlen",
         "Starke F&E, Markenkompetenz, Innovationsfähigkeit, Qualitätsmanagement"],
        ["Fokus (Nische)",
         "Konzentration auf ein Segment mit entweder Kosten- oder Differenzierungsvorteil",
         "Tiefes Segmentwissen, spezifische Kundennähe, begrenzte Ressourcen effizient eingesetzt"],
    ],
    cap="Tabelle T.3: Generische Wettbewerbsstrategien nach Porter. Quelle: Eigene Darstellung in Anlehnung an Porter (1985, [S. X]).",
)
para(
    "Die Analyse von Wettbewerbsintensität und Differenzierungsmerkmalen (Abschnitt "
    "3.3.4) baut auf diesem Framework auf: Sie fragt, welche generische Strategie "
    "das Unternehmen verfolgt, wie klar diese positioniert ist und wo gegenüber "
    "den Wettbewerbern tatsächlich überlegene Merkmale bestehen."
)

# ═══════════════════════════════════════════════════════════════════════════
# RESOURCE-BASED VIEW
# ═══════════════════════════════════════════════════════════════════════════
divider()
h2("Resource-Based View (RBV) und Kernkompetenzen")

para(
    "Der Resource-Based View (RBV) stellt das Gegenstück zu Porters marktorientiertem "
    "Ansatz dar: Während Porter Wettbewerbsvorteile primär aus der Branchenstruktur "
    "ableitet, erklärt der RBV sie durch die einzigartigen Ressourcen und Fähigkeiten "
    "des einzelnen Unternehmens."
)

h3("Herkunft und theoretische Entwicklung")
para(
    "Die intellektuellen Wurzeln des RBV liegen in Edith Penroses Werk "
    "„The Theory of the Growth of the Firm“ (1959), in dem sie das "
    "Unternehmen erstmals als Bündel produktiver Ressourcen konzipierte, nicht "
    "als bloße Produktionseinheit. Birger Wernerfelt formalisierte diesen Ansatz "
    "1984 unter dem Begriff Resource-Based View. Die für das strategische Management "
    "entscheidende Ausformulierung stammt von Jay Barney, der 1991 im "
    "„Journal of Management“ die präzisen Bedingungen definierte, unter "
    "denen Ressourcen nachhaltige Wettbewerbsvorteile begründen können "
    "(Barney, 1991, S. 99–100). Die Kernthese: Unternehmen derselben Branche "
    "unterscheiden sich systematisch in ihrer Ressourcenausstattung, und diese "
    "Heterogenität ist die eigentliche Quelle von Leistungsunterschieden – "
    "nicht allein die Marktstruktur (Barney, 1991, S. 101)."
)

h3("Das VRIN-Framework")
para(
    "Barney (1991, S. 105–112) operationalisiert die Bedingung des nachhaltigen "
    "Wettbewerbsvorteils durch vier Kriterien, die eine Ressource kumulativ erfüllen muss:"
)
tbl(
    ["Kriterium", "Bedeutung", "Analytische Frage"],
    [
        ["Valuable\n(wertvoll)",
         "Die Ressource muss dem Unternehmen ermöglichen, eine Chance zu nutzen oder eine Bedrohung zu neutralisieren",
         "Trägt die Ressource zur Wertschöpfung oder Kostenreduktion bei?"],
        ["Rare\n(selten)",
         "Die Ressource darf nicht bei einer Vielzahl von Wettbewerbern vorhanden sein",
         "Verfügen viele oder wenige Wettbewerber über dieselbe Ressource?"],
        ["Inimitable\n(nicht imitierbar)",
         "Die Ressource darf für Wettbewerber nicht oder nur mit prohibitiv hohem Aufwand kopierbar sein",
         "Gibt es Imitationsbarrieren? (historische Pfadabhängigkeit, kausale Ambiguität, soziale Komplexität)"],
        ["Non-substitutable\n(nicht substituierbar)",
         "Es darf keine strategisch gleichwertige, aber anders geartete Ressource existieren",
         "Kann der Wettbewerbsvorteil dieser Ressource durch eine andere Art von Ressource repliziert werden?"],
    ],
    cap="Tabelle T.4: VRIN-Kriterien nach Barney (1991). Quelle: Eigene Darstellung in Anlehnung an Barney (1991, S. 105–112).",
)
para(
    "Barney unterscheidet vier Ressourcenkategorien: finanzielle, physische, "
    "humane und organisationale (immaterielle) Ressourcen. In der Analyse werden "
    "die Ressourcen des Unternehmens zunächst systematisch erfasst und anschließend "
    "gegen die VRIN-Kriterien geprüft, um strategisch relevante von lediglich "
    "notwendigen Ressourcen zu trennen."
)

h3("Kernkompetenzen nach Prahalad & Hamel")
para(
    "Prahalad und Hamel (1990) ergänzen den RBV um eine dynamisch-organisationale "
    "Perspektive. In ihrem Beitrag „The Core Competence of the Corporation“ "
    "(Harvard Business Review, 1990) argumentieren sie, dass nachhaltiger "
    "Wettbewerbsvorteil nicht aus einzelnen Ressourcen, sondern aus dem kollektiven "
    "Lernen der Organisation entsteht – insbesondere der Fähigkeit, verschiedene "
    "Technologien und Produktionskompetenzen zu integrieren und zu koordinieren "
    "(Prahalad & Hamel, 1990, S. 82)."
)
para(
    "Eine Kernkompetenz im Sinne von Prahalad und Hamel erfüllt drei Testkriterien "
    "(Prahalad & Hamel, 1990, S. 83–84):"
)
tbl(
    ["Kriterium", "Erläuterung"],
    [
        ["Marktzugang",
         "Die Kompetenz ermöglicht Zugang zu einer Vielzahl unterschiedlicher Märkte – sie ist keine rein produktspezifische Fähigkeit."],
        ["Wahrgenommener Kundenwert",
         "Sie leistet einen signifikanten Beitrag zum wahrgenommenen Nutzen des Endprodukts aus Kundensicht."],
        ["Schwere Imitierbarkeit",
         "Sie ist für Wettbewerber schwer nachzuahmen, weil sie auf kollektivem Wissen, Erfahrung und organisationalen Routinen basiert."],
    ],
    cap="Tabelle T.5: Testkriterien für Kernkompetenzen nach Prahalad & Hamel (1990, S. 83–84). Quelle: Eigene Darstellung.",
)

# ═══════════════════════════════════════════════════════════════════════════
# WERTKETTENANALYSE
# ═══════════════════════════════════════════════════════════════════════════
divider()
h2("Wertkettenanalyse nach Porter")

para(
    "Die Wertkettenanalyse wurde von Porter (1985) in „Competitive Advantage: "
    "Creating and Sustaining Superior Performance“ entwickelt. Sie bildet die "
    "operative Konkretisierung zur Frage, wo genau im Unternehmen Wettbewerbsvorteile "
    "entstehen oder Kostennachteile bestehen. Während das Five-Forces-Modell die "
    "Branchenebene analysiert, richtet die Wertkette den Blick nach innen und "
    "dekonstruiert das Unternehmen in seine strategisch relevanten Aktivitäten "
    "(Porter, 1985, [S. X])."
)
para(
    "Das Grundprinzip: Jedes Unternehmen ist eine Abfolge von Aktivitäten, die "
    "zusammen den Wert erzeugen, den ein Kunde für ein Produkt oder eine "
    "Dienstleistung zu zahlen bereit ist. Der Gewinn ist die Differenz zwischen "
    "diesem erzeugten Gesamtwert und den Kosten, die bei allen Aktivitäten "
    "zusammen anfallen (Porter, 1985, [S. X]). Wettbewerbsvorteile entstehen "
    "dort, wo ein Unternehmen einzelne Aktivitäten effizienter oder "
    "differenzierter ausführt als seine Konkurrenten."
)

h3("Struktur der Wertkette")
para(
    "Porter (1985, [S. X]) unterscheidet zwischen primären Aktivitäten, die "
    "direkt zur Leistungserstellung und -vermarktung beitragen, und "
    "Unterstützungsaktivitäten, die die Primäraktivitäten ermöglichen und "
    "untereinander vernetzen."
)
tbl(
    ["Kategorie", "Aktivität", "Inhalt"],
    [
        ["Primäre\nAktivitäten", "Eingangslogistik",
         "Beschaffung, Lagerung und Verteilung der Eingangsressourcen; Lieferantenmanagement"],
        ["", "Operationen",
         "Transformation der Eingangsressourcen in das Endprodukt (Fertigung, Montage, Qualitätsprüfung)"],
        ["", "Ausgangslogistik",
         "Lagerung und Distribution des fertigen Produkts zum Kunden; Auftragsabwicklung"],
        ["", "Marketing & Vertrieb",
         "Preisfindung, Kommunikation, Vertriebskanäle, Kundenakquise"],
        ["", "Service",
         "Installation, Inbetriebnahme, Wartung, Reparatur, Schulung nach dem Verkauf"],
        ["Unterstützungs-\naktivitäten", "Unternehmensinfrastruktur",
         "Allgemeines Management, Finanzen, Recht, Qualitätsmanagement, IT"],
        ["", "Personalwirtschaft",
         "Rekrutierung, Auswahl, Entwicklung, Vergütung aller Mitarbeitenden"],
        ["", "Technologieentwicklung",
         "F&E, Prozessinnovation, Produktentwicklung, Know-how-Management"],
        ["", "Beschaffung",
         "Strategisches Einkaufsmanagement für alle Ressourcen (nicht nur Rohstoffe)"],
    ],
    cap="Tabelle T.6: Struktur der Wertkette nach Porter (1985). Quelle: Eigene Darstellung in Anlehnung an Porter (1985, [S. X]).",
)

h3("Auswertungslogik")
para(
    "In der Auswertung werden für jede Aktivität drei Fragen gestellt: (1) Welche "
    "Kosten entstehen hier? (2) Welchen Beitrag leistet die Aktivität zur "
    "Differenzierung aus Kundensicht? (3) Bestehen Verknüpfungen ("
    "„Linkages“) zu anderen Aktivitäten, die Kostensenkungen oder "
    "Qualitätsverbesserungen ermöglichen? Gerade an diesen Schnittstellen zwischen "
    "Aktivitäten entstehen häufig die wesentlichsten Quellen von "
    "Wettbewerbsvorteil, da sie von Konkurrenten besonders schwer zu erkennen "
    "und zu imitieren sind (Porter, 1985, [S. X])."
)

# ═══════════════════════════════════════════════════════════════════════════
# SWOT
# ═══════════════════════════════════════════════════════════════════════════
divider()
h2("SWOT-Analyse")

para(
    "Die SWOT-Analyse nimmt unter den strategischen Analyseinstrumenten eine "
    "besondere Stellung ein: Sie ist kein eigenständiges Erhebungswerkzeug, "
    "sondern ein integratives Syntheseinstrument, das die Ergebnisse aller "
    "vorangegangenen Analysen zu einer strategischen Gesamtbeurteilung "
    "zusammenführt (Whittington et al., 2024, [S. X])."
)

h3("Herkunft")
para(
    "Die SWOT-Analyse entstammt der Strategielehre an der Harvard Business School "
    "der 1960er-Jahre. Kenneth Andrews formalisierte das Konzept 1971 in "
    "„The Concept of Corporate Strategy“ und etablierte damit das "
    "Grundprinzip der Synthese aus interner Stärken-/Schwächenanalyse und "
    "externer Chancen-/Risikoanalyse. Das Akronym SWOT steht für Strengths, "
    "Weaknesses, Opportunities und Threats (Whittington et al., 2024, [S. X])."
)

h3("Konzeptionelle Logik und Auswertung")
para(
    "Die methodische Stärke der SWOT liegt nicht in der isolierten Betrachtung "
    "der vier Felder, sondern in ihrer Kreuzanalyse. Aus der Kombination interner "
    "Faktoren (Stärken, Schwächen) mit externen Faktoren (Chancen, Risiken) "
    "lassen sich vier Normstrategietypen ableiten, die strategische Handlungs-"
    "optionen strukturieren (Whittington et al., 2024, [S. X]):"
)
tbl(
    ["Kombination", "Strategietyp", "Grundlogik"],
    [
        ["Stärken × Chancen (SO)",
         "Ausbaustrategie",
         "Vorhandene Stärken gezielt einsetzen, um externe Chancen zu erschließen und Wachstum zu generieren."],
        ["Stärken × Risiken (ST)",
         "Verteidigungsstrategie",
         "Stärken mobilisieren, um externe Bedrohungen abzuwehren oder deren Wirkung zu minimieren."],
        ["Schwächen × Chancen (WO)",
         "Aufholstrategie",
         "Interne Defizite abbauen, um externe Chancen nutzen zu können, die bislang unzugänglich sind."],
        ["Schwächen × Risiken (WT)",
         "Rückzugsstrategie",
         "Risikominimierung durch gleichzeitigen Abbau von Schwächen und Vermeidung von Risikokombinationen."],
    ],
    cap="Tabelle T.7: SWOT-Normstrategien. Quelle: Eigene Darstellung in Anlehnung an Whittington et al. (2024, [S. X]).",
)
para(
    "Methodisch ist zu beachten, dass die SWOT-Analyse nur so valide ist wie die "
    "Qualität der Voranalysen, auf die sie aufbaut: PESTEL und Five Forces liefern "
    "die externen Faktoren (O und T), während RBV und Wertkettenanalyse die "
    "internen Faktoren (S und W) begründen. Eine SWOT ohne fundierte Voranaly-"
    "sen reduziert sich auf subjektive Auflistungen ohne analytischen Mehrwert "
    "(Whittington et al., 2024, [S. X])."
)

# ═══════════════════════════════════════════════════════════════════════════
# LITERATURVERZEICHNIS
# ═══════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("Literaturverzeichnis")

entries = [
    ("Barney, J. (1991). Firm resources and sustained competitive advantage. ",
     "Journal of Management",
     ", 17(1), 99–120. https://doi.org/10.1177/014920639101700108"),
    ("Grant, R. M. (2024). ",
     "Contemporary strategy analysis",
     " (12. Aufl.). Wiley. ISBN 978-1-394-25159-9"),
    ("Porter, M. E. (1980). ",
     "Competitive strategy: Techniques for analyzing industries and competitors.",
     " The Free Press. ISBN 978-0-02-925360-1"),
    ("Porter, M. E. (1985). ",
     "Competitive advantage: Creating and sustaining superior performance.",
     " The Free Press. ISBN 978-0-02-925090-7"),
    ("Porter, M. E. (2008). The five competitive forces that shape strategy. ",
     "Harvard Business Review",
     ", 86(1), 78–93."),
    ("Prahalad, C. K., & Hamel, G. (1990). The core competence of the corporation. ",
     "Harvard Business Review",
     ", 68(3), 79–91."),
    ("Whittington, R., Régner, P., Angwin, D., Johnson, G., & Scholes, K. (2024). ",
     "Exploring strategy: Text and cases",
     " (13. Aufl.). Pearson Education. ISBN 978-1-292-47953-8"),
]
for b, i, a in entries:
    bib_entry(b, i, a)

# ── save ────────────────────────────────────────────────────────────────
out = '/home/user/test-data1/Kapitel_3_Theoretische_Grundlagen.docx'
doc.save(out)
print(f"Gespeichert: {out}")
